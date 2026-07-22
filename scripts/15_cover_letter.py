"""Generate cover letter for Applied Energy"""
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54); s.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

def p(text, bold=False, after=6, align=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = bold
    para.paragraph_format.space_after = Pt(after)
    if align: para.alignment = align
    return para

p('', after=24)
p('Date: July 2026', after=12)
p('', after=12)
p('Dear Editor-in-Chief,', after=12)
p('Applied Energy', after=6)
p('Elsevier', after=18)

p('Re: Submission of Original Research Article', bold=True, after=18)

p('We are pleased to submit our manuscript entitled "Data-Driven Nuclear Baseload and Battery '
  'Storage Dispatch Optimization under Duck Curve Uncertainty: A Conformalized Machine Learning '
  'Approach" for consideration as a Full Length Research Article in Applied Energy.', after=12)

p('This paper addresses the optimal co-dispatch of nuclear baseload generation and battery '
  'energy storage under the net load uncertainty created by high solar penetration (the "duck '
  'curve"). Using real 2023 California ISO (CAISO) operational data, we develop an integrated '
  'pipeline from probabilistic forecasting to robust dispatch optimization: an XGBoost day-ahead '
  'net load forecast (R² = 0.854, +18.5% skill over persistence) feeds a conformalized quantile '
  'regression (CQR) uncertainty layer into a budget-of-uncertainty robust MILP dispatch model, '
  'benchmarked against a two-stage stochastic programming baseline under an identical information '
  'set. Removing 2.26 GW of nuclear capacity raises system cost by 13.2% and weekly CO₂ emissions '
  'by 17.3%; a 50-week annual sweep, together with targeted robustness checks against realized '
  'hydro generation and CCGT fleet aggregation granularity, shows this economic premium is joined '
  'by a distinct, larger reliability-value regime during summer capacity-stress weeks. Robust '
  'dispatch itself is shown to be a cost–robustness–emissions tradeoff, making forecast quality '
  'an emissions-abatement lever with direct carbon-policy implications.', after=12)

p('We believe this work fits squarely within the scope of Applied Energy: it combines '
  'energy-related forecasting and decision-making, optimization methods for energy systems, '
  'and energy policy and economic analysis, using real operational data and a fully reproducible '
  'methodology. This manuscript has not been published previously and is not under consideration '
  'for publication elsewhere. The author has approved the manuscript and its submission to '
  'Applied Energy.', after=12)

p('Thank you for considering our submission. We look forward to your response.', after=18)

p('Sincerely,', after=6)
p('Cagatay Kuban', after=3)
p('Independent Researcher, Ankara, Türkiye', after=3)
p('cagataykuban@gmail.com', after=3)

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'paper')
path = os.path.join(OUT, 'cover_letter.docx')
doc.save(path)
print(f'Saved: {path}')
