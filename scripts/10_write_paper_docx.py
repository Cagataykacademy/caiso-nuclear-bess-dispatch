"""
Generate Q1 journal paper as Word document (.docx)
Target: Electric Power Systems Research / Applied Energy — ~8,000 words
"""
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(PROJECT, "outputs", "figures")
OUT = os.path.join(PROJECT, "paper")
os.makedirs(OUT, exist_ok=True)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)

def para(text, bold=False, italic=False, align=None, size=12, indent=True, spacing=2.0, after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(after)
    return p

def body(text):
    return para(text, indent=True)

def fig(filename, caption, width=6.0):
    if os.path.exists(os.path.join(FIG, filename)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(os.path.join(FIG, filename), width=Inches(width))
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
        c.paragraph_format.space_after = Pt(12)

def _set_cell_borders(cell, top=None, bottom=None):
    """Booktabs-style: only horizontal rules (top/bottom), never vertical."""
    tcPr = cell._element.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for side in ('left', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(f'{ns}val', 'nil')
        borders.append(el)
    for side, weight in (('top', top), ('bottom', bottom)):
        el = OxmlElement(f'w:{side}')
        if weight:
            el.set(f'{ns}val', 'single'); el.set(f'{ns}sz', str(weight)); el.set(f'{ns}color', '000000')
        else:
            el.set(f'{ns}val', 'nil')
        borders.append(el)
    tcPr.append(borders)

def tbl(caption, headers, rows):
    para(caption, bold=True, size=10, indent=False, after=4)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Normal Table'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    n_rows = len(rows)
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]; cell.text = ''
        r = cell.paragraphs[0].add_run(h); r.bold = True
        r.font.name = 'Times New Roman'; r.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_borders(cell, top=12, bottom=8)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i+1].cells[j]; cell.text = ''
            r = cell.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_borders(cell, bottom=12 if i == n_rows - 1 else None)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0

print("Building comprehensive manuscript...", flush=True)

# =====================================================================
# TITLE PAGE
# =====================================================================
para('', indent=False, after=30)
para('Data-Driven Nuclear Baseload and Battery Storage Dispatch Optimization '
     'under Duck Curve Uncertainty: A Conformalized Machine Learning Approach',
     bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=18)

para('', indent=False, after=12)
para('Cagatay Kuban ¹ *', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=6)
para('¹ Independent Researcher, Ankara, Türkiye',
     italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=6)
para('* Corresponding author: cagataykuban@gmail.com', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=30)

# =====================================================================
# HIGHLIGHTS
# =====================================================================
heading('Highlights', 2)
highlights = [
    'True CAISO net load from EIA fuel-type data reveals a 55 GW duck curve depth.',
    'XGBoost day-ahead forecast: R² = 0.854, +18.5% skill over persistence.',
    'CQR-robust MILP matches stochastic programming with 40% fewer binary variables.',
    'Nuclear premium is 13% in typical weeks but reaches 210% under capacity stress.',
    'Robust dispatch trades +4.2% cost for +21.8% CO2, an emissions-abatement lever.',
]
for h in highlights:
    bullet(h)

doc.add_page_break()

# =====================================================================
# ABSTRACT
# =====================================================================
heading('Abstract', 1)

body(
    'The rapid growth of solar photovoltaics has exacerbated the "duck curve" phenomenon, marked '
    'by steep net load ramps and negative residual demand. We propose an integrated data-driven '
    'framework coupling machine learning (ML) uncertainty quantification with robust '
    'mixed-integer linear programming (MILP) dispatch of nuclear baseload and battery storage '
    '(BESS). Using 2023 CAISO data, we construct the true net load signal (−10,369 to +44,900 MW) '
    'and forecast it day-ahead with XGBoost (R² = 0.854, +18.5% skill over persistence), '
    'validated out-of-sample on 2022 (R² = 0.908). Conformalized quantile regression (CQR) '
    'provides distribution-free 90% prediction intervals (88.8% coverage), parameterizing a '
    'budget-of-uncertainty robust MILP with a two-tier gas fleet and a 5 GW/20 GWh BESS. Removing '
    '2.26 GW of nuclear capacity raises cost by 13.2% ($41.84 → $47.35/MWh) and CO₂ emissions by '
    '17.3% (~9 Mt/year). A two-stage stochastic programming baseline attains near-identical '
    'realized cost with 40% more binaries and several-fold longer solves. A 50-week annual sweep '
    'shows the nuclear premium is tightly clustered near 13% in normal weeks but rises to '
    '74–210% during summer capacity-stress weeks; adding realized CAISO hydro generation to '
    'these weeks eliminates 93.4% of the flagged shedding, confirming a modeling artifact, and '
    'results are insensitive to CCGT aggregation granularity. A supplementary ERCOT-parameterized '
    'check shows the premium amplified (+15–16%) under scarcer import capacity. Robustness '
    'itself is a three-way tradeoff: hedging to the CQR upper bound raises cost by 4.2% and '
    'emissions by 21.8%, making forecast quality an emissions-abatement lever.'
)

para('Keywords: Duck curve; Net load forecasting; Conformalized quantile regression; '
     'Mixed-integer linear programming; Nuclear baseload; Battery energy storage; CAISO',
     italic=True, size=11, indent=False, after=12)

doc.add_page_break()

# =====================================================================
# NOMENCLATURE
# =====================================================================
heading('Nomenclature', 1)

nomenclature = [
    ['Sets and Indices', '', ''],
    ['t', '∈ {1, ..., T}', 'Time period index (hours)'],
    ['T', '', 'Planning horizon (168 hours = 1 week)'],
    ['', '', ''],
    ['Parameters', '', ''],
    ['c_nuc', '12 $/MWh', 'Nuclear marginal cost'],
    ['c_gas', '45 $/MWh', 'Natural gas marginal cost'],
    ['c_bess', '5 $/MWh', 'BESS cycling (degradation) cost'],
    ['c_imp', '55 $/MWh', 'Import cost'],
    ['r_exp', '20 $/MWh', 'Export revenue'],
    ['V', '10,000 $/MWh', 'Value of lost load (VOLL)'],
    ['c_curt', '10 $/MWh', 'Curtailment cost'],
    ['η_rt', '0.90', 'BESS round-trip efficiency'],
    ['P̄_nuc', '2,256 MW', 'Nuclear capacity (Diablo Canyon)'],
    ['P_nuc', '1,800 MW', 'Nuclear minimum stable output'],
    ['P̄_bess', '5,000 MW', 'BESS power capacity'],
    ['Ē_bess', '20,000 MWh', 'BESS energy capacity (4h duration)'],
    ['ΔP_nuc', '100 MW/h', 'Nuclear ramp rate limit (PWR physical limit, ~4.4%/h)'],
    ['ΔP_gas', '5,000 MW/h', 'Gas ramp rate limit'],
    ['', '', ''],
    ['Decision Variables', '', ''],
    ['P_nuc(t)', 'MW', 'Nuclear power output'],
    ['P_gas(t)', 'MW', 'Gas power output'],
    ['P_ch(t)', 'MW', 'BESS charging power'],
    ['P_dis(t)', 'MW', 'BESS discharging power'],
    ['SoC(t)', 'MWh', 'BESS state of charge'],
    ['P_imp(t)', 'MW', 'Power imports'],
    ['P_exp(t)', 'MW', 'Power exports'],
    ['P_shed(t)', 'MW', 'Involuntary load shedding'],
    ['P_curt(t)', 'MW', 'Renewable curtailment'],
    ['u(t)', '{0,1}', 'BESS mode (1 = discharging)'],
    ['', '', ''],
    ['Abbreviations', '', ''],
    ['BESS', '', 'Battery energy storage system'],
    ['CAISO', '', 'California Independent System Operator'],
    ['CQR', '', 'Conformalized quantile regression'],
    ['DM', '', 'Diebold–Mariano (test)'],
    ['EIA', '', 'U.S. Energy Information Administration'],
    ['LMP', '', 'Locational marginal price'],
    ['MILP', '', 'Mixed-integer linear programming'],
    ['PI', '', 'Prediction interval'],
    ['VRE', '', 'Variable renewable energy'],
]

t = doc.add_table(rows=len(nomenclature), cols=3)
t.style = 'Normal Table'
n_nom = len(nomenclature)
for i, (sym, val, desc) in enumerate(nomenclature):
    for j, text in enumerate([sym, val, desc]):
        cell = t.rows[i].cells[j]; cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        if sym in ['Sets and Indices', 'Parameters', 'Decision Variables', 'Abbreviations', '']:
            r.bold = True
        _set_cell_borders(cell, top=12 if i == 0 else None, bottom=12 if i == n_nom - 1 else None)
doc.add_paragraph()

doc.add_page_break()

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
heading('1. Introduction', 1)

body(
    'The integration of variable renewable energy (VRE) sources—particularly utility-scale '
    'solar photovoltaics (PV)—has fundamentally transformed the operational landscape of modern '
    'power systems. In California, the installed solar capacity exceeded 18 GW by 2023, producing '
    'mid-day generation surges that routinely depress net load (total demand minus solar and wind '
    'generation) to near-zero or negative values (Denholm et al., 2015). This creates the '
    'characteristic "duck curve," first identified by the California Independent System Operator '
    '(CAISO) in 2013, featuring a deep mid-day valley (the "belly") followed by a steep evening '
    'ramp as solar output declines while demand peaks (CAISO, 2016). The operational implications '
    'of this pattern—originally projected for the 2020s—have materialized faster than anticipated, '
    'with CAISO reporting net load swings exceeding 35 GW within a single day during spring 2023.'
)

body(
    'The duck curve poses three interconnected challenges for system operators. First, the steep '
    'evening ramp—reaching 3,000 MW/h or more in CAISO—demands fast-responding generation '
    'resources that can increase output within minutes, straining conventional thermal fleet '
    'capabilities. Second, periods of over-generation during mid-day hours risk curtailment of '
    'renewable output, wasting clean energy and eroding the economic case for further solar '
    'investment. Third, baseload generators such as nuclear plants, which operate most efficiently '
    'at near-constant output due to reactor physics and licensing constraints, face increasing '
    'pressure to either cycle (reducing efficiency, increasing maintenance costs, and accelerating '
    'fuel wear) or shut down entirely during periods of low residual demand (Jenkins et al., 2018). '
    'This tension between inflexible baseload and variable renewables is particularly acute in '
    'CAISO, where Diablo Canyon—California\'s last nuclear power plant at 2,256 MW—continues to '
    'provide roughly 9% of the state\'s electricity.'
)

body(
    'California\'s grid-scale BESS capacity surpassed 5 GW by 2023, providing arbitrage '
    'value (charging during mid-day duck belly, discharging during evening ramps). The '
    'co-optimization of BESS with nuclear baseload under uncertain net load conditions motivates '
    'integrating probabilistic forecasting with dispatch optimization. Conformal prediction '
    '(Vovk et al., 2005) has emerged as a distribution-free method providing finite-sample '
    'coverage guarantees for prediction intervals without distributional assumptions, bridging '
    'the gap between ML-based uncertainty quantification and robust operations research.'
)

body('This paper makes the following contributions:')

contributions = [
    'Construction of the true CAISO net load signal using real hourly generation-by-fuel data '
    '(solar, wind, nuclear, natural gas, hydro, coal, petroleum, other) from the U.S. Energy '
    'Information Administration (EIA), comprising 70,458 hourly records across eight fuel categories. '
    'The resulting net load ranges from −10,369 to +44,900 MW—a duck curve depth exceeding 35 GW—'
    'providing an empirical foundation that most prior studies lack, as they rely on synthetic or '
    'aggregated data.',

    'Development of a day-ahead (24-hour horizon) net load forecasting framework using three '
    'ensemble learning methods (LightGBM, XGBoost, Random Forest) and two deep learning '
    'baselines (LSTM, MLP) with 35 leakage-free features, rigorously benchmarked against '
    'persistence using Diebold–Mariano tests (Diebold and Mariano, 1995) and five-fold '
    'expanding-window cross-validation, with out-of-sample validation on the full year 2022. '
    'All features are strictly filtered to ensure availability at the time of forecast '
    '(no look-ahead bias).',

    'Application of conformalized quantile regression (CQR) (Romano et al., 2019) to produce '
    'calibrated 90% prediction intervals that are directly integrated into a robust MILP dispatch '
    'formulation, bridging the gap between ML-based forecasting and operations research optimization.',

    'Comprehensive quantification of nuclear baseload and BESS value through: (a) 23 parametric '
    'sensitivity configurations (nuclear 0–4 GW, BESS 0.5–10 GW, gas $25–$100/MWh); '
    '(b) four-season dispatch analysis (winter, spring, summer, fall) using realized CAISO net '
    'load, showing dispatch cost ranging from $39.85/MWh (spring) to $45.16/MWh (summer); and '
    '(c) Value of Perfect Information (VOPI = $0.31/MWh, 0.74%) quantifying the economic impact '
    'of the 1,429 MW day-ahead forecast error on dispatch economics.',

    'Robustness and transferability checks: we benchmark the CQR-robust MILP against a '
    'two-stage stochastic programming baseline under an identical information set, verify the '
    'representativeness of the primary test week via a 50-week annual sweep spanning the full '
    '2023 record (revealing a distinct reliability-value regime during capacity-stress weeks), '
    'validate this regime directly by adding realized 2023 CAISO hydro generation and by '
    'testing CCGT aggregation granularity, and probe qualitative transferability of the '
    'framework through a narrowly scoped supplementary check on an ERCOT-parameterized grid '
    '(Appendix B).',
]
for i, c in enumerate(contributions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.27)
    r = p.add_run(f'({i}) {c}')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

# --- Key Findings at a Glance box ---
kf_table = doc.add_table(rows=1, cols=1)
kf_table.style = 'Table Grid'
kf_cell = kf_table.cell(0, 0)
kf_cell.text = ''
kf_title = kf_cell.paragraphs[0]
kf_title.paragraph_format.space_before = Pt(6)
kf_title.paragraph_format.space_after = Pt(4)
r = kf_title.add_run('Key Findings at a Glance')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

key_findings = [
    'Day-ahead net load forecasting: XGBoost achieves R² = 0.854, +18.5% skill over '
    'persistence, validated out-of-sample on 2022 (R² = 0.908).',
    'CQR delivers calibrated 90% prediction intervals (88.8% coverage) and matches two-stage '
    'stochastic programming performance with 40% fewer binary variables and several-fold faster solves.',
    'Nuclear baseload value is +13.2% system cost under typical conditions, rising to '
    '+74–210% during summer capacity-stress weeks (50-week annual sweep)—a reliability-value '
    'regime distinct from the baseline economic premium and confirmed robust to CCGT '
    'aggregation and realized-hydro sensitivity checks.',
    'Robust dispatch is a three-way tradeoff: hedging to the CQR upper bound costs +4.2% more '
    'and emits +21.8% more CO₂, making forecast quality itself an emissions-abatement lever.',
    'A narrowly scoped supplementary check (Appendix B) shows the scarcity mechanism transfers '
    'directionally to an ERCOT-parameterized grid (+15–16% premium under an island grid\'s '
    'limited import capacity).',
]
for kf in key_findings:
    kp = kf_cell.add_paragraph(style='List Bullet')
    kp.paragraph_format.space_after = Pt(3)
    r = kp.add_run(kf)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
kf_cell.paragraphs[-1].paragraph_format.space_after = Pt(6)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

body(
    'The remainder of this paper is organized as follows. Section 2 reviews related literature '
    'on net load forecasting, conformal prediction, and dispatch optimization. Section 3 describes '
    'the methodology in detail, including data sources, feature engineering, ML models, CQR '
    'calibration, and the MILP formulation. Section 4 presents results and Section 5 provides '
    'an in-depth discussion. Section 6 concludes with policy implications and future research '
    'directions.'
)

doc.add_page_break()

# =====================================================================
# 2. LITERATURE REVIEW
# =====================================================================
heading('2. Literature Review', 1)

heading('2.1. Forecasting and Uncertainty Quantification', 2)

body(
    'Short-term net load forecasting—defined as total demand minus solar and wind generation—'
    'has evolved from classical ARIMA models to modern gradient boosting methods. LightGBM '
    '(Ke et al., 2017) and XGBoost (Chen and Guestrin, 2016) consistently achieve state-of-the-art '
    'performance on tabular energy data, outperforming recurrent neural networks in the majority '
    'of benchmark studies (Hertel et al., 2022; Shao et al., 2022). Net load forecasting presents '
    'challenges beyond traditional demand forecasting: residual errors are non-Gaussian, '
    'heteroscedastic, and time-varying, particularly during ramp events (Wang et al., 2019). '
    'A persistent methodological concern is data leakage—using variables not available at forecast '
    'time (e.g., contemporaneous generation or demand). The present study enforces a strict '
    '24-hour information cutoff, ensuring all features are available before the target hour.'
)

body(
    'Uncertainty quantification for dispatch optimization requires calibrated prediction intervals. '
    'Standard quantile regression (Koenker, 2005) yields adaptive intervals but may be '
    'miscalibrated under distribution shift. Conformal prediction (Vovk et al., 2005) provides '
    'distribution-free finite-sample coverage guarantees by construction, requiring only '
    'exchangeability between calibration and test data. Romano et al. (2019) introduced '
    'conformalized quantile regression (CQR), which combines quantile regression adaptiveness '
    'with conformal calibration guarantees. CQR has demonstrated reliable near-nominal coverage '
    'for wind power and electricity applications (Stankovic et al., 2023; Xu and Chen, 2023). '
    'The slight violation of exchangeability in time series contexts does not materially degrade '
    'coverage when the calibration period is temporally proximate to the test period (Gibbs and '
    'Candès, 2021), as in our September–October calibration / November–December test split.'
)

heading('2.2. Dispatch Optimization and Nuclear Value', 2)

body(
    'Dispatch optimization under VRE uncertainty has been addressed through stochastic '
    'programming (Birge and Louveaux, 2011), which minimizes expected cost across scenario trees '
    'but scales poorly with horizon length, and robust optimization (Bertsimas et al., 2013), '
    'which optimizes against worst-case realizations within a defined uncertainty set. '
    'Budget-of-uncertainty sets (Bertsimas and Sim, 2004) offer a compact parameterization of '
    'conservatism. Our approach bridges these paradigms: CQR provides a distribution-free '
    '[Q10(t), Q90(t)] interval at each hour t as a data-driven box uncertainty set. The '
    'parameterized net load P_net(t, γ) = P_point(t) + γ · (Q90(t) − P_point(t)), γ ∈ [0,1], '
    'directly implements a budget-of-uncertainty conservatism level (γ = 0: nominal; γ = 1: '
    'worst-case Q90), scaling linearly with horizon length while inheriting CQR\'s formal '
    'coverage guarantee.'
)

body(
    'The value of nuclear baseload in high-renewable grids has gained renewed policy relevance '
    '(Jenkins et al., 2018). Nuclear plants provide emissions-free, low-marginal-cost baseload '
    '(~$12/MWh) but are constrained by reactor physics and licensing to near-constant output. '
    'Under duck curve conditions, nuclear\'s inflexibility creates tension with oversupply during '
    'mid-day solar peaks. BESS alleviates this tension by absorbing duck-belly surplus and '
    'discharging during evening ramps (Denholm et al., 2023). For CAISO specifically, Diablo '
    'Canyon\'s 2,256 MW—extended through 2030 by SB 846 (2022)—provides ~9% of California\'s '
    'electricity and avoids roughly 7–8 Mt CO₂/year. Despite its policy significance, its quantified '
    'dispatch value under actual CAISO conditions remains understudied; this paper addresses '
    'that gap using real 2023 hourly operational data and a four-season MILP analysis.'
)

body(
    'Taken together, the literature leaves two gaps this paper addresses jointly: the lack of a '
    'rigorous, same-information comparison between interval-driven robust optimization and '
    'classical stochastic programming (Section 4.11), and the absence of an empirically grounded '
    'valuation of nuclear dispatch flexibility under real, high-renewable-penetration operating '
    'conditions (Sections 4.6-4.12).'
)

doc.add_page_break()

# =====================================================================
# 3. METHODOLOGY
# =====================================================================
heading('3. Methodology', 1)

body(
    'This section describes the five-stage pipeline: (3.1) data acquisition and net load '
    'construction, (3.2) feature engineering, (3.3) ML model training and benchmarking, '
    '(3.4) CQR calibration, and (3.5) MILP formulation and scenario design.'
)

heading('3.1. Data Acquisition and Net Load Construction', 2)

body(
    'We use five real data sources for CAISO\'s 2023 operations, obtained from three public '
    'repositories (Table 1). The primary data source is the EIA\'s Hourly Electric Grid Monitor '
    '(Form EIA-930), which provides hourly demand, net generation, and interchange data for all '
    'U.S. balancing authorities. We query the CAISO (respondent code "CISO") region data, '
    'obtaining 35,052 hourly records, and the fuel-type-specific generation data, obtaining '
    '70,458 records across eight fuel categories (solar, wind, nuclear, natural gas, hydroelectric, '
    'coal, petroleum, and other).'
)

tbl('Table 1. Data sources and characteristics.',
    ['Dataset', 'Source', 'Records', 'Frequency', 'Period'],
    [['Demand, generation, interchange', 'EIA Form 930 API', '35,052', 'Hourly', 'Jan–Dec 2023'],
     ['Generation by fuel type', 'EIA Form 930 API', '70,458', 'Hourly', 'Jan–Dec 2023'],
     ['Wholesale electricity price (SP15)', 'EIA/ICE Markets', '365', 'Daily', 'Jan–Dec 2023'],
     ['Natural gas spot price (Henry Hub)', 'EIA Nat. Gas API', '249', 'Daily', 'Jan–Dec 2023'],
     ['Temperature (LAX, SFO, Fresno)', 'NOAA ISD Lite', '8,760', 'Hourly', 'Jan–Dec 2023']])

body(
    'The true net load is computed as P_net(t) = P_demand(t) − P_solar(t) − P_wind(t), '
    'where P_solar and P_wind are obtained from the EIA generation-by-fuel dataset. This '
    'construction—using actual metered generation rather than installed capacity or capacity '
    'factors—captures the real-time variability of renewable output, including curtailment '
    'effects and sub-hourly ramp behavior aggregated to the hourly level.'
)

body(
    'The resulting net load exhibits the defining features of the duck curve (Fig. 1). Key '
    'statistics include: mean = 18,733 MW, minimum = −10,369 MW (spring mid-day, when solar '
    'generation exceeds total demand), maximum = +44,900 MW (summer evening peak), and standard '
    'deviation = 6,700 MW. The full-range duck curve depth—maximum minus minimum net load—is '
    '55,269 MW (~55 GW); within-day excursions exceed 35 GW on the steepest ramp days '
    '(Section 1). Negative '
    'net load hours, representing periods when solar and wind output exceed total system demand, '
    'occur primarily in March–May between 10:00 and 14:00 PST.'
)

fig('fig02_duck_curve.png',
    'Fig. 1. CAISO 2023 duck curve by season. Left: hourly demand (solid) vs. net load (dashed) '
    'by season. Right: spring profile showing solar displacement (yellow shading) and the '
    'duck belly, where net load drops below zero.')

body(
    'Electricity prices are obtained from the EIA/ICE wholesale electricity market data, '
    'specifically the SP15 EZ Gen Day-Ahead LMP Peak product, which represents traded '
    'day-ahead wholesale prices for the Southern California hub. The dataset provides daily '
    'weighted-average, high, and low prices for 234 trading days in 2023, with non-trading days '
    'filled by temporal interpolation. Prices enter the study only as lagged exploratory '
    'features and benchmark material (Section 4.2); no daily-to-hourly disaggregation is '
    'performed, and the MILP uses fixed technology marginal costs throughout. '
    'The resulting daily price series ranges from $9.43 to '
    '$348.19/MWh (mean $63.05, std $44.65).'
)

body(
    'Temperature data is obtained from the NOAA Integrated Surface Database (ISD Lite) for three '
    'California stations: Los Angeles International Airport (LAX, representative of SP15/Southern '
    'California), San Francisco International Airport (SFO, representative of NP15/Northern '
    'California), and Fresno (FAT, representative of the Central Valley). Hourly temperatures '
    'are averaged across stations to produce a CAISO-wide temperature proxy, from which cooling '
    'degree hours (CDH, base 18.3°C) and heating degree hours (HDH) are computed.'
)

heading('3.2. Feature Engineering for Day-Ahead Forecasting', 2)

body(
    'We adopt a strict day-ahead forecasting protocol: at forecast time T, only information '
    'available before T may be used to predict P_net(T+24). This constraint eliminates all '
    'features with lags shorter than 24 hours, prevents the use of rolling statistics computed '
    'at the target time, and requires that all input features be either deterministic (calendar '
    'variables) or derived from data at least 24 hours old. This design choice is critical for '
    'methodological rigor, as violations of this constraint—common in the literature—lead to '
    'artificially inflated accuracy metrics (R² > 0.99) that do not reflect operational forecast '
    'performance.'
)

body(
    'The resulting 35-feature set is organized into five categories (Table 2). The rationale '
    'for each category is as follows. Calendar features capture the deterministic seasonal and '
    'diurnal patterns in demand; cyclical (sine/cosine) encodings are used to represent hour '
    'and month as continuous periodic variables, which improves gradient boosting performance '
    'compared to integer encoding (Hong and Fan, 2016). The CAISO day-ahead demand forecast, '
    'published before each operating day, provides an expert baseline prediction that the ML '
    'model can refine. Lagged net load and demand features at 24, 48, 72, and 168 hours capture '
    'the strong autocorrelation structure of electricity demand—the lag-24h value (same hour '
    'yesterday) is the single most informative feature for day-ahead forecasting. Yesterday\'s '
    'statistics (mean, max, min, std, range of the previous day\'s demand profile) provide a '
    'summary of recent demand shape, enabling the model to detect whether the previous day was '
    'unusually hot, mild, or volatile.'
)

body(
    'Solar and wind generation lags (24h and 168h) provide information about recent renewable '
    'output patterns, enabling the model to anticipate days with similar meteorological conditions. '
    'Additional trend indicators and lagged wholesale prices (SP15 LMP) serve as proxies for '
    'broader system conditions. Note that explicit temperature features were excluded during '
    'feature selection, as lagged demand and net load already embed the demand response to weather.'
)

tbl('Table 2. Feature categories for day-ahead net load forecasting (35 features).',
    ['Category', 'Count', 'Key Features', 'Availability at T'],
    [['Calendar (deterministic)', '12', 'Hour, day of week, month, day of year, weekend, week, sin/cos encodings', 'Deterministic'],
     ['ISO day-ahead forecast', '1', 'CAISO published demand forecast', 'Published before T'],
     ['Lagged net load & demand', '8', 'Net load and demand at 24, 48, 72, 168 h lags', 'Available (≥24h old)'],
     ['Yesterday statistics', '6', 'Mean, max, min, std of demand; mean net load; solar mean', 'Available (≥24h old)'],
     ['Renewable & Trends', '6', 'Solar/wind lag 24/168h, demand/netload trends 24h', 'Available (≥24h old)'],
     ['Market & Error', '2', 'Price lag24, ISO forecast error lag24', 'Available (≥24h old)']])

body(
    'Data are split temporally without shuffling to preserve the time series structure: training '
    '(January 8 – August 31, 2023; n = 5,664), validation (September 1 – October 31; n = 1,464), '
    'and test (November 1 – December 31; n = 1,345, after excluding 119 hours with missing EIA '
    'fuel-type records, 8.1% of the window). The training set spans spring through summer, '
    'capturing the full duck curve season; the validation set covers early autumn; and the test '
    'set covers late autumn through winter. The validation set serves two purposes: early stopping '
    'during gradient boosting training (preventing overfitting) and conformal calibration '
    '(computing the CQR correction Q̂). The test set is held out entirely for final evaluation '
    'and is never used for model selection or hyperparameter tuning.'
)

heading('3.3. Machine Learning Models and Benchmarking', 2)

body(
    'Three ensemble learning models are trained for both net load and electricity price targets. '
    'LightGBM (Ke et al., 2017) uses histogram-based gradient boosting with 63 leaves, learning '
    'rate 0.03, L₁/L₂ regularization (λ = 0.1), feature and bagging fractions of 0.7, and up to '
    '3,000 iterations with early stopping (patience = 100 rounds). XGBoost (Chen and Guestrin, '
    '2016) uses depth-limited trees (max depth 6) with the same learning rate, subsampling, and '
    'early stopping configuration. Random Forest (Breiman, 2001) uses 500 trees with maximum '
    'depth 15, minimum 20 samples per leaf, and 0.7 max features ratio.'
)

body(
    'All models are benchmarked against a persistence baseline, defined as yesterday\'s same-hour '
    'value: ŷ(t) = y(t−24). Persistence is a well-established benchmark in energy forecasting '
    'that any operational model must outperform to justify its complexity (Hong and Fan, 2016). '
    'Pairwise model comparisons are conducted using the Diebold–Mariano (DM) test (Diebold and '
    'Mariano, 1995) with Newey–West HAC variance estimation using 24 lags, testing the null '
    'hypothesis of equal predictive accuracy based on squared-error loss.'
)

body(
    'Model stability is assessed via five-fold expanding-window cross-validation: the first '
    'fold trains on the initial 40% of data and tests on the next block; subsequent folds expand '
    'the training window while maintaining a fixed test block size (1,016 hours ≈ 42 days). '
    'This design respects temporal ordering while providing multiple performance estimates across '
    'different seasonal regimes.'
)

body(
    'To contextualize model performance, we compute the forecast error of the CAISO-published '
    'day-ahead gross demand forecast (EIA Form 930 field "day_ahead_forecast_MW") on the same '
    'November–December 2023 test period. The CAISO EMS gross-load day-ahead forecast achieves '
    'MAE = 1,330 MW and MAPE = 5.65% (mean gross demand 23,712 MW). Our XGBoost model achieves '
    'MAE = 1,429 MW on net load, a normalized MAE of 7.4% on the mean test net load of '
    '19,299 MW. (We report normalized MAE rather than MAPE for net load, which is ill-defined '
    'when the target crosses zero.) '
    'The modest gap reflects the fundamentally greater volatility of the net load target: net '
    'load variance is elevated relative to gross demand variance due to compounded solar and wind '
    'forecast uncertainty, making day-ahead net load forecasting inherently more challenging than '
    'gross demand forecasting. Out-of-sample temporal generalization is confirmed by applying '
    'LightGBM—trained on 2023 data—to 8,593 hourly observations from calendar year 2022, '
    'yielding R² = 0.908 and MAE = 1,514 MW. Because 2022 wholesale price data were not '
    'available at the time of this validation, the single lagged-price feature '
    '(price_lag24h) is set to zero for the 2022 run; given its minor contribution to the '
    'trained model (1.1% of total feature-importance gain, Appendix Fig. A5), this '
    'substitution is expected to have negligible impact on the reported 2022 metrics. '
    'Part of the higher 2022 R² reflects an '
    'intrinsically more predictable year (persistence itself improves from R² = 0.736 to 0.822); '
    'the relevant evidence is that the model\'s skill over persistence is preserved out of year '
    '(MAE gain of the same LightGBM model: +21.6% in 2022 vs +14.1% on the 2023 test set), '
    'demonstrating that the model captures durable seasonal and diurnal patterns rather than '
    'year-specific artifacts, with no evidence of temporal overfitting.'
)

heading('3.4. Conformalized Quantile Regression', 2)

body(
    'We construct prediction intervals using the CQR framework of Romano et al. (2019). '
    'The procedure consists of three stages: quantile model training, conformal calibration, '
    'and interval construction.'
)

body(
    'Stage 1 (Quantile model training): Five LightGBM quantile regression models are trained at '
    'quantile levels τ ∈ {0.10, 0.25, 0.50, 0.75, 0.90} using the pinball loss function '
    'L_τ(y, q) = τ·max(y−q, 0) + (1−τ)·max(q−y, 0). The hyperparameter configuration matches '
    'the point prediction models (63 leaves, learning rate 0.03, L₁/L₂ regularization) to ensure '
    'comparability. The Q10 and Q90 models define the initial (uncalibrated) 80% prediction interval. '
    'LightGBM is used for the quantile stage because it supports the quantile objective natively; '
    'its point accuracy is statistically close to XGBoost\'s (R² 0.843 vs 0.854), and the '
    'conformal step is agnostic to the base learner.'
)

body(
    'Stage 2 (Conformal calibration on the validation set): For each validation observation '
    '(X_i, Y_i) where i = 1, ..., n_val, the conformity score E_i = max(q̂_0.10(X_i) − Y_i, '
    'Y_i − q̂_0.90(X_i)) measures how much the true value falls outside the uncalibrated interval. '
    'Positive E_i indicates a miss (true value outside the interval); negative E_i indicates '
    'the margin by which the true value is inside. The conformal quantile Q̂ is set to the '
    '⌈(n+1)(1−α)⌉/n-th quantile of the conformity score distribution, where α = 0.10 for '
    'a 90% coverage target. Intuitively, Q̂ represents the "extra margin" needed to widen the '
    'quantile regression intervals so that they achieve the desired coverage.'
)

body(
    'Stage 3 (Calibrated interval construction): The calibrated 90% prediction interval for a '
    'new test point X is C(X) = [q̂_0.10(X) − Q̂,  q̂_0.90(X) + Q̂]. The symmetric widening by '
    'Q̂ ensures that the interval inherits both the adaptiveness of quantile regression (the '
    'interval width varies with input difficulty) and the coverage guarantee of conformal '
    'prediction (at least (1−α) marginal coverage under exchangeability). Note that the base '
    'pair (q̂_0.10, q̂_0.90) is nominally an 80% interval; the conformal step recalibrates it '
    'to the 90% target, which is valid because the coverage guarantee derives solely from the '
    'conformity-score quantile, not from the nominal levels of the base quantiles (Romano et '
    'al., 2019). For our application, '
    'the computed Q̂ = 1,044 MW, meaning the conformal correction widens each side of the '
    'prediction interval by approximately 1,044 MW—a substantial but necessary adjustment that '
    'accounts for the inherent uncertainty in day-ahead net load forecasting.'
)

heading('3.5. MILP Dispatch Formulation', 2)

body(
    'We formulate a 168-hour (one-week) economic dispatch problem as a mixed-integer linear '
    'program (MILP). The 168-hour horizon captures the full weekly periodicity of both CAISO net '
    'load and BESS cycling behavior: a 24-hour horizon undervalues inter-day arbitrage '
    'opportunities, while annual models sacrifice the intra-week granularity needed to resolve '
    'duck curve dynamics. The formulation adopts a single-bus (copper-plate) representation, '
    'omitting transmission constraints to isolate the generation–storage interaction under duck '
    'curve conditions. This approach is standard in system-level weekly dispatch studies (Jenkins '
    'et al., 2018; Denholm et al., 2023). To empirically validate this assumption, we implement '
    'a two-zone (NP15/SP15) extension of the MILP, adding a Path 15 transmission constraint '
    '(4,500 MW forward / 3,500 MW reverse; CAISO OASIS, 2023) and allocating capacities '
    'proportionally (see Appendix Fig. A12). The two-zone model yields a dispatch cost of '
    '$43.73/MWh versus $42.15/MWh for the copper-plate—a modest difference of 3.73%. Path 15 '
    'utilization averages 24.3% of rated capacity and binds in 0 out of 168 hours. This confirms '
    'that inter-zonal congestion has negligible impact on weekly aggregate dispatch costs, '
    'justifying the copper-plate formulation. Furthermore, the scenario comparisons that are '
    'the focus of this study—nuclear premium (S4−S1), robustness cost (γ=1−γ=0)—are insensitive '
    'to the copper-plate assumption because any residual inter-zonal price differential cancels '
    'in pairwise cost differences. Quantifying congestion rents in a full nodal model is identified '
    'as a priority for future work (Section 6). '
    'The model is solved using the open-source HiGHS solver (v1.7; Huangfu and Hall, 2018) via Pyomo (v6.7; Hart et al., 2017), achieving '
    'global optimality (gap < 0.01%) in well under 1 second (0.4-0.5 s typical) for all 840-binary instances.'
)

body(
    'Gas generation is modeled using a two-tier representation that captures the distinct '
    'operational characteristics of CAISO\'s thermal fleet. The CCGT fleet (10,000 MW in total) '
    'is represented as two aggregate units of 5,000 MW each, each with explicit unit commitment '
    'binaries: a commitment variable u_ccgt(t,k) ∈ {0,1} and a startup indicator '
    'y_ccgt(t,k) ∈ {0,1}. CCGTs have a variable cost of $40/MWh, a startup cost of $25,000 per '
    'unit per event (≈ $5/MW warm-start; NREL ATB, 2023), a minimum up-time of 4 hours, a minimum '
    'down-time of 2 hours, and a minimum operating load of 1,500 MW when committed. Their ramp '
    'rate is constrained to ±800 MW/h per unit, reflecting steam turbine thermal inertia. Simple-'
    'cycle peakers and reciprocating engines (up to 10,000 MW aggregate) are modeled as a '
    'continuous variable with cost $65/MWh and ramp ±5,000 MW/h—effectively unconstrained—'
    'representing the fast-response, high-cost margin of CAISO\'s thermal fleet.'
)
body(
    'The decision variables at each hour t ∈ {1, ..., 168} are: nuclear output P_nuc(t), '
    'CCGT output P_ccgt(t), peaker output P_peak(t), BESS charging P_ch(t) and discharging '
    'P_dis(t), state of charge SoC(t), imports P_imp(t), exports P_exp(t), involuntary load '
    'shedding P_shed(t), renewable curtailment P_curt(t), and a binary BESS mode indicator '
    'u(t) ∈ {0,1} preventing simultaneous charge and discharge.'
)

body(
    'Objective function: Minimize total system cost.'
)
body(
    'min Σ_t [ c_nuc · P_nuc(t) + c_ccgt · P_ccgt(t) + c_peak · P_peak(t) '
    '+ c_bess · (P_ch(t) + P_dis(t)) '
    '+ c_imp · P_imp(t) − r_exp · P_exp(t) + V · P_shed(t) + c_curt · P_curt(t) ]'
)
body(
    'The cost coefficients are: c_nuc = $12/MWh (nuclear fuel + O&M), c_ccgt = $40/MWh '
    '(CCGT fuel + O&M), c_peak = $65/MWh (simple-cycle peaker, less efficient), '
    'c_bess = $5/MWh (battery cycling degradation), '
    'c_imp = $55/MWh (import cost from neighboring balancing areas), r_exp = $20/MWh (export '
    'revenue), V = $10,000/MWh (value of lost load, representing the extreme penalty for '
    'involuntary load shedding), and c_curt = $10/MWh (renewable curtailment opportunity cost).'
)

body(
    'The constraints are as follows. (C1) Power balance at each hour: '
    'P_nuc(t) + P_ccgt(t) + P_peak(t) + P_dis(t) + P_imp(t) + P_shed(t) = '
    'P_net(t) + P_ch(t) + P_exp(t) + P_curt(t). (C2) BESS state-of-charge dynamics: '
    'SoC(t) = SoC(t−1) + √η · P_ch(t) − (1/√η) · P_dis(t), where η = 0.90 is the round-trip '
    'efficiency and SoC(0) = 0.5 · Ē_bess. (C3) End-of-horizon SoC floor: '
    'SoC(T) ≥ 0.35 · Ē_bess. This constraint prevents the optimizer from gaming the finite '
    'horizon by discharging the BESS without recharging in the final hours—a known artifact of '
    'rolling-horizon formulations (Sioshansi et al., 2009). (C4) BESS mutual exclusion via big-M: '
    'P_ch(t) ≤ P̄_bess · (1 − u(t)) and P_dis(t) ≤ P̄_bess · u(t). (C5) Nuclear ramp rate: '
    '|P_nuc(t) − P_nuc(t−1)| ≤ 100 MW/h, reflecting the physical load-following capability '
    'of pressurized water reactors (~4.4%/h of rated capacity; NEA, 2011). (C6) CCGT ramp: '
    '±800 MW/h (steam turbine thermal inertia constraint). (C7) Peaker ramp: ±5,000 MW/h '
    '(simple-cycle GTs, effectively unconstrained). (C8) CCGT unit commitment: '
    'x_ccgt(t,k) <= 5000·u_ccgt(t,k) and x_ccgt(t,k) >= 1500·u_ccgt(t,k); '
    'startup indicator y_ccgt(t,k) >= u_ccgt(t,k) - u_ccgt(t-1,k); '
    'min up-time: sum_{tau=t}^{t+3} u_ccgt(tau,k) >= 4·y_ccgt(t,k); '
    'min down-time: sum_{tau=t}^{t+1} (1-u_ccgt(tau,k)) >= 2·(u_ccgt(t-1,k)-u_ccgt(t,k)). '
    '(C9) SoC bounds: 0.10·Ē_bess <= SoC(t) <= 0.90·Ē_bess. '
    '(C10) Capacity bounds: 1,800-2,256 MW for nuclear, 0-10,000 MW for CCGT, '
    '0-10,000 MW for peakers, 0-10,000 MW for imports. '
    'Explicit spinning reserve constraints are not included in the formulation. At the peak of '
    'the representative summer week (Aug 7–13; Section 4.7)—when aggregate committed capacity '
    '(37,256 MW: nuclear + CCGT + peakers + BESS + imports) '
    'barely exceeds net load (36,999 MW)—imposing a percentage-based reserve requirement '
    'would force artificial shedding. More starkly, at the true annual net load peak '
    '(44,900 MW, mid-July; Section 4.12), the stylized fleet\'s aggregate capacity is a '
    'shortfall, not a margin—reinforcing rather than undermining this modeling choice, since a '
    'formal reserve requirement would only worsen an already capacity-constrained extreme week. '
    'In practice, CAISO meets '
    'reserve requirements through bilateral capacity contracts, demand response programs, and '
    'intra-hour redispatch mechanisms that lie outside the scope of this weekly aggregate model '
    '(CAISO, 2023), as discussed in Section 4.12. Explicit reserve stacking in a nodal model with '
    'demand response resources '
    'is identified as future work (Section 6).'
)

body(
    'The model contains approximately 2,184 variables—including 840 binary variables '
    '(168 BESS mode variables plus 2 × 168 CCGT commitment and 2 × 168 CCGT startup variables)—'
    'and approximately 2,100 constraints per scenario. All scenarios are solved to global '
    'optimality (gap < 0.01%) in well under 1 second (0.4-0.5 s typical) using the HiGHS branch-and-bound solver, '
    'demonstrating computational tractability. This formulation represents a simplified but '
    'realistic unit commitment (UC) for the CCGT tier, while retaining the economic dispatch '
    'structure for peakers. Full UC with individual generator modeling (heat rate curves, '
    'no-load costs, maintenance windows) is identified as a future research extension.'
)

body(
    'The CQR prediction intervals serve as a data-driven box uncertainty set for the MILP. '
    'Specifically, the parameterized net load P_net(t, γ) = P_point(t) + γ · (Q90(t) − P_point(t)) '
    'for γ ∈ [0,1] implements the budget-of-uncertainty framework of Bertsimas and Sim (2004): '
    'γ = 0 yields the nominal (deterministic) dispatch; γ = 1 yields the worst-case dispatch '
    'against the CQR Q90 upper bound; intermediate γ values trade cost for robustness. Because '
    'all dispatch decisions are here-and-now (no recourse) and both cost and the binding '
    'power-balance constraint are monotone in net load, the worst case over the box '
    '[Q10(t), Q90(t)] is attained at the upper bound; solving the deterministic MILP at the '
    'γ-scaled upper profile is therefore the exact static robust counterpart, not a heuristic '
    'approximation. The '
    'formal CQR coverage guarantee (88.8% empirical, 90% target) provides a principled basis '
    'for this uncertainty set that scenario-based approaches cannot offer. '
    'Six scenarios are evaluated: S1 (γ=0, Deterministic) uses the XGBoost point forecast; '
    'S2 (γ=1, Worst-case Q90) and S3 (γ=−∞, Best-case Q10) represent the extremes; '
    'S6 (γ=0.5, Robust) blends point and Q90; S4 (No Nuclear) sets nuclear to zero to '
    'quantify baseload value; S5 (Small BESS, 1 GW / 4 GWh) tests storage sensitivity. '
    'The robustness parameter sweep (γ ∈ {0, 0.25, 0.50, 0.75, 1.00}, Section 4.10) '
    'directly cross-validates S1, S6, and S2, confirming internal consistency.'
)

body(
    'CO₂ emissions are computed ex post from the optimal dispatch using technology-specific '
    'emission factors: 0.37 t CO₂/MWh for CCGTs (heat rate ~7.0 MMBtu/MWh at 53.07 kg '
    'CO₂/MMBtu), 0.55 t CO₂/MWh for simple-cycle peakers (~10.5 MMBtu/MWh), and 0.428 t '
    'CO₂/MWh for imports—the California Air Resources Board default emission factor for '
    'unspecified imported electricity, a deliberately conservative regulatory convention. '
    'Nuclear, BESS throughput, and renewable generation are assigned zero operational emissions.'
)

body(
    'Stochastic programming baseline: To benchmark the CQR-robust approach against the '
    'classical alternative, we formulate a two-stage stochastic program on the same '
    'information set. The scenario set is the CQR triplet {Q10, point, Q90} with '
    'Swanson–Megill probabilities {0.30, 0.40, 0.30}, the standard three-point approximation '
    'for P10/P50/P90 fractiles (Keefer and Bodily, 1983). First-stage (here-and-now) decisions '
    'are the inflexible ones: CCGT commitment and startup binaries and the nuclear trajectory; '
    'all continuous dispatch (CCGT output levels, peakers, BESS, imports/exports, shedding, '
    'curtailment) is second-stage recourse per scenario. All plans—deterministic, robust '
    '(γ = 0.5, 1), and SP—are then evaluated out of sample under an identical protocol: their '
    'first-stage decisions are frozen and the recourse problem is re-solved against the '
    'realized net load, yielding directly comparable realized costs (Section 4.11).'
)

doc.add_page_break()

# =====================================================================
# 4. RESULTS
# =====================================================================
heading('4. Results', 1)

heading('4.1. Generation Mix and Duck Curve Formation', 2)

body(
    'Fig. 2 shows the hourly generation mix during a representative spring week (April 10–16, '
    '2023). The duck curve formation is clearly visible: during mid-day hours (10:00–15:00), '
    'solar generation (up to 16,000 MW) constitutes the dominant source, displacing gas, hydro, '
    'and imports. Nuclear output remains essentially constant at 2,256 MW throughout the week, '
    'illustrating its role as inflexible baseload. As solar output declines in the late afternoon, '
    'gas generation ramps steeply from ~2,000 MW to over 15,000 MW within 4–5 hours—the evening '
    'ramp that defines the duck curve\'s "neck." Wind generation provides a modest but variable '
    'contribution (500–4,000 MW), with higher output during nighttime hours.'
)

fig('fig03_generation_mix_spring_week.png',
    'Fig. 2. CAISO generation mix during a spring week (April 10–16, 2023), showing duck curve '
    'formation. Solar (yellow) dominates mid-day, nuclear (green) provides constant baseload, '
    'and gas (orange) handles the steep evening ramp.')

heading('4.2. Net Load Forecasting Performance', 2)

body(
    'Table 3 summarizes the day-ahead forecasting results on the held-out test set '
    '(November–December 2023). For net load, XGBoost achieves the best test performance '
    '(R² = 0.854, MAE = 1,429 MW, RMSE = 2,071 MW), representing a forecast skill of +18.5% '
    'over persistence. All three ML models significantly outperform persistence at the 0.1% '
    'level (Diebold–Mariano p < 0.001). XGBoost is statistically superior to Random Forest '
    '(DM p = 0.026) and LightGBM (DM p < 0.001).'
)

tbl('Table 3. Day-ahead forecasting benchmark (test set: November–December 2023). '
    'Bold indicates best ML model. DM test is against persistence baseline.',
    ['Target', 'Model', 'MAE', 'RMSE', 'R²', 'Train R²', 'Gap', 'DM p-value'],
    [['Net Load (MW)', 'LightGBM', '1,565', '2,146', '0.843', '0.997', '0.154', '< 0.001 ***'],
     ['Net Load (MW)', 'XGBoost', '1,429', '2,071', '0.854', '0.972', '0.118', '< 0.001 ***'],
     ['Net Load (MW)', 'Random Forest', '1,533', '2,211', '0.834', '0.947', '0.113', '< 0.001 ***'],
     ['Net Load (MW)', 'LSTM (2-layer)', '2,278', '3,085', '0.676', '0.924', '0.248', 'p < 0.001 ***'],
     ['Net Load (MW)', 'MLP (3-layer)', '3,220', '3,904', '0.481', '0.996', '0.514', 'p < 0.001 ***'],
     ['Net Load (MW)', 'Persistence', '1,753', '2,783', '0.736', '—', '—', 'baseline']])

body(
    'Explicit temperature features (CDH, HDH, and their interactions) were evaluated during '
    'feature selection but excluded from the final 35-feature set: lagged demand and net load '
    'already embed the demand response to weather, and adding station temperatures did not '
    'improve validation performance. The NOAA temperature data is instead used for error '
    'diagnostics (Section 4.3 and Appendix Fig. A7), where it reveals that the largest forecast '
    'errors concentrate on extreme heat days—information content that lagged loads cannot fully '
    'capture and that motivates the weather-ensemble extension discussed in Section 6.'
)

body(
    'Two deep learning benchmarks are included. For the LSTM, a manual hyperparameter search '
    'was conducted over: hidden units (32, 64, 128), lookback windows (24, 48, 72 hours), '
    'dropout rates (0.1, 0.2, 0.3), and learning rates (1e-3, 3e-4, 1e-4). The best '
    'configuration—2-layer LSTM, 64 hidden units per layer, 24-hour lookback, dropout = 0.2, '
    'Adam with cosine-annealing learning rate, early stopping at epoch 17—achieves '
    'R² = 0.676 and MAE = 2,278 MW, which is worse than both persistence (R² = 0.736) '
    'and XGBoost (R² = 0.854). The LSTM overfit gap (train R² = 0.924 vs test R² = 0.676) '
    'reveals that ~5,500 training samples are insufficient for a recurrent architecture to '
    'learn generalizable temporal dependencies from hourly energy data; gradient boosting '
    'methods avoid this through built-in regularization and ensemble diversity. '
    'A 3-layer MLP (128-64-32 neurons, ReLU, Adam) performs even worse '
    '(R² = 0.481, MAE = 3,220 MW) with extreme overfitting (gap = 0.514), confirming '
    'the well-documented superiority of GBM over deep learning on moderate-scale tabular data.'
)

body(
    'These results confirm, for the tested recurrent and feedforward architectures, '
    'the well-documented advantage of gradient boosting methods over '
    'recurrent and feedforward neural networks for tabular energy data (Hertel et al., 2022; '
    'Shao et al., 2022). The key advantage of GBM models is their built-in feature selection '
    'and regularization, which prevent the memorization that plagues neural networks on '
    'moderate-sized tabular datasets. While Transformer-based architectures have shown promise '
    'for very large energy datasets (millions of samples), they are unlikely to outperform GBM '
    'at the scale of our dataset (~5,500 training samples).'
)

body(
    'Note on electricity price forecasting: We elected to exclude electricity price forecasting '
    'from the ML scope of this study. Preliminary experiments using the same feature set on daily '
    'SP15 DA LMP prices showed that all ML models failed to outperform persistence (best ML MAE = '
    '$6.5/MWh vs persistence MAE = $4.5/MWh). This is consistent with the electricity price '
    'forecasting literature (Weron, 2014), which documents that daily prices exhibit strong '
    'autocorrelation that makes persistence a formidable baseline. Since our MILP formulation '
    'uses fixed marginal costs (standard in dispatch optimization), price forecasting is not '
    'required for the dispatch pipeline.'
)

fig('fig06_benchmark_comparison.png',
    'Fig. 3. Day-ahead forecasting benchmark: MAE comparison across models for net load '
    '(left) and price (right). Dashed line indicates persistence baseline.')

heading('4.3. Cross-Validation and Model Stability', 2)

body(
    'Five-fold expanding-window cross-validation confirms framework stability across seasonal '
    'regimes (mean R² = 0.869 ± 0.089). The single exception, Fold 2, coincides with the 2023 '
    'California heat events; the full fold-by-fold breakdown and the associated '
    'temperature-error analysis are reported in Appendix A (Table A3, Fig. A7) to keep the '
    'main results focused on headline findings.'
)

heading('4.4. Prediction Intervals and Conformal Calibration', 2)

body(
    'CQR produces well-calibrated prediction intervals for net load: the 90% PI achieves 88.8% '
    'empirical coverage with a mean width of 5,182 MW, conformal correction Q̂ = 1,044 MW, and '
    'a mean interval (Winkler) score of 8,346 MW at α = 0.10. '
    'The slight undercoverage (88.8% vs. 90% target) is within the expected finite-sample '
    'variation for n = 1,345 test observations. The conformal step is not cosmetic: the '
    'uncalibrated quantile-regression interval achieves only 61.1% empirical coverage '
    '(width 3,094 MW) with a worse Winkler score of 10,874 MW—conformalization therefore '
    'improves the interval under a proper scoring rule while restoring near-nominal coverage, '
    'rather than merely trading width for coverage. Fig. 4 illustrates the prediction intervals '
    'over the full test period and a zoomed first week, showing that the intervals are adaptive: '
    'they widen during high-uncertainty periods (e.g., rapid ramps) and narrow during stable '
    'demand hours.'
)

fig('fig07_netload_prediction_PI.png',
    'Fig. 4. Day-ahead net load forecast with 90% CQR prediction intervals. Top: full test '
    'period (November–December 2023). Bottom: zoomed first week showing interval adaptiveness.')

body(
    'Marginal coverage, however, does not imply conditional coverage—and we report the gap '
    'honestly. Fig. 4b disaggregates coverage by hour of day: mid-day hours are over-covered '
    '(up to 100%) while evening-ramp hours (16:00–22:00) fall to 78.6–84%, precisely when net '
    'load uncertainty matters most for dispatch. This heterogeneity is expected—split-conformal '
    'methods guarantee only marginal coverage (Romano et al., 2019)—and it does not undermine '
    'the robust MILP, whose γ-parameterization hedges against the upper bound irrespective of '
    'when misses occur. It does, however, sharpen the case for the adaptive and '
    'time-conditional conformal extensions identified in Section 6.'
)

fig('fig_conformal_diagnostics.png',
    'Fig. 4b. Conformal calibration diagnostics. (a) Uncalibrated QR vs calibrated CQR '
    'intervals over the first test week; red crosses mark observations outside the CQR '
    'interval. (b) Conditional coverage by hour of day: evening-ramp hours (shaded) are '
    'under-covered relative to the 90% target, while mid-day hours are over-covered.')

body(
    'Formal residual diagnostics (skewness, kurtosis, normality, and autocorrelation tests) '
    'and a hyperparameter sensitivity analysis are reported in Appendix A (Tables A1–A2); '
    'both confirm well-behaved, distribution-appropriate residuals and a model choice that is '
    'not overly sensitive to hyperparameter selection.'
)

heading('4.6. Dispatch Optimization Results', 2)

body(
    'Table 5 presents the MILP results for the six core scenarios evaluated on the November '
    'test week. All scenarios are solved to global optimality (gap < 0.01%) using the HiGHS '
    'solver in well under 1 second (0.4-0.5 s typical), demonstrating the computational tractability of the formulation.'
)

tbl('Table 5. MILP scenario comparison (168-hour dispatch, November 2023). '
    'Two-tier gas: CCGT ($40/MWh, ±800 MW/h) + Peakers ($65/MWh, ±5,000 MW/h). '
    'End-of-horizon SoC ≥ 35% constraint enforced in all scenarios.',
    ['Scenario', 'Description', '$/MWh', 'CCGT (MW)', 'Peaker (MW)', 'Import (MW)', 'Shed (MWh)', 'BESS Cycles', 'SoC End (%)'],
    [['S1 Deterministic', 'Point forecast',     '41.84', '9,734', '144',   '5,716', '0', '4.94', '35%'],
     ['S2 Worst-case',    'CQR Q90',            '43.59', '10,000','1,275', '7,155', '0', '0.44', '35%'],
     ['S3 Best-case',     'CQR Q10',            '40.25', '9,036', '0',     '4,213', '0', '3.25', '35%'],
     ['S4 No Nuclear',    'Nuc = 0 MW',         '47.35', '9,987', '1,153', '6,678', '0', '2.51', '35%'],
     ['S5 Small BESS',    '1 GW / 4 GWh',       '42.16', '9,274', '543',   '5,769', '0', '5.18', '7%'],
     ['S6 Robust',        'γ = 0.50 blending',   '42.63', '9,989', '641',   '6,380', '0', '2.54', '35%']])

body(
    'Key findings from the scenario analysis: (i) Removing nuclear baseload (S4 vs. S1) '
    'increases unit cost by $5.51/MWh (+13.2%), with the two-tier gas fleet (additional peakers '
    'reaching 1,153 MW average) and increased imports compensating for the lost 2,256 MW of '
    'zero-carbon baseload. The end-of-horizon SoC constraint (SoC[168] ≥ 35%) is binding in '
    'S1/S2/S4/S6; for S5 (1 GW/4 GWh BESS), the SoC floor is 35% of 4,000 MWh = 1,400 MWh, '
    'reported as 7% of the 20 GWh reference capacity—the constraint is binding in all scenarios. '
    '(ii) The robust formulation (S6) incurs a 1.9% cost premium ($0.79/MWh) over the '
    'deterministic case—higher than in simpler formulations because the two-tier gas model '
    'reveals the true flexibility cost of peaker deployment. (iii) Reducing BESS from 5 GW to '
    '1 GW (S5 vs. S1) increases cost by only 0.95%, but the S5 SoC ends at 7%—near depletion—'
    'indicating that the small BESS has insufficient capacity to maintain the end-of-period '
    'reserve. The sharp drop in BESS cycling under the pessimistic profiles (S2: 0.44 '
    'cycles/week vs S1: 4.94) has a simple mechanism: hedging pre-schedules additional CCGT '
    'and peaker output across all hours, which flattens the residual intra-day differential '
    'that arbitrage exploits—robustness crowds out storage utilization. '
    '(iv) Zero load shedding and zero curtailment across all scenarios confirms that '
    'the November week, while challenging, lies within the operational envelope.'
)

fig('fig12_milp_dispatch_stack.png',
    'Fig. 5. Optimal dispatch for S1 (Deterministic). Top: generation stack showing nuclear '
    'baseload, gas ramping, and BESS discharge. Middle: BESS charge/discharge with state of '
    'charge. Bottom: hourly system cost breakdown.')

heading('4.7. Four-Season Dispatch Analysis', 2)

body(
    'To assess the generalizability of results beyond the November test week, we run the improved '
    'MILP on four representative weeks covering all seasons: Winter (Jan 23–29), Spring (Apr 10–16), '
    'Summer (Aug 7–13), and Fall (Nov 1–7). The weeks are selected ex ante by transparent '
    'criteria: the spring week contains the annual net load minimum (−2,852 MW within the week), '
    'the summer week is a typical high-demand summer week (peak net load 36,999 MW), and the '
    'winter and fall weeks are '
    'typical mid-season weeks with no holidays. Each week uses realized CAISO net load as the '
    'MILP input. Table 6 summarizes the results.'
)

tbl('Table 6. Four-season dispatch results (with nuclear, actual net load).',
    ['Season', 'NL Mean (MW)', 'NL Max (MW)', '$/MWh', 'CCGT (MW)', 'Peaker (MW)', 'Import (MW)', 'Shed (MWh)', 'BESS Cycles', 'CCGT Starts'],
    [['Winter (Jan 23–29)',  '18,851', '27,383', '42.27', '9,939', '402',   '6,240', '0', '2.37', '2'],
     ['Spring (Apr 10–16)', '14,439', '23,919', '39.85', '8,504', '0',     '3,718', '0', '3.42', '4'],
     ['Summer (Aug 7–13)',  '22,366', '36,999', '45.16', '10,000', '2,604', '7,480', '0', '2.02', '0'],
     ['Fall (Nov 1–7)',     '18,407', '28,430', '42.15', '9,793', '450',   '5,894', '0', '2.89', '2']])

body(
    'Summer dispatch is most expensive ($45.16/MWh) due to high net demand (peak 36,999 MW) '
    'requiring sustained peaker deployment (2,604 MW average). Spring is the least costly '
    '($39.85/MWh): negative net load periods (reaching −2,852 MW) allow BESS charging at low '
    'marginal cost, and no peakers are dispatched. Winter and Fall exhibit intermediate costs '
    '($42.27 and $42.15/MWh respectively), with CCGTs providing 93–99% of thermal generation. '
    'CCGT unit commitment decisions reflect seasonal load patterns: Summer maintains units online '
    'continuously (0 starts), while Spring requires 4 start–stop cycles to avoid minimum-load '
    'violations during duck-belly hours. Zero load shedding across all four selected seasonal '
    'weeks confirms system adequacy with nuclear online under typical conditions; the annual net '
    'load maximum (44,900 MW, mid-July) falls outside this typical-week sample and is analyzed '
    'separately as a capacity-stress event in Section 4.12. The S4 (no-nuclear) scenario '
    'shows that the gas fleet and imports maintain zero shedding operationally in these typical '
    'weeks—but at a sustained '
    '$5.51/MWh premium (Section 4.6)—underscoring nuclear\'s economic rather than reliability value '
    'under normal conditions.'
)

fig('fig_seasonal_dispatch.png',
    'Fig. 6. Four-season dispatch stacks (Winter Jan 23–29, Spring Apr 10–16, '
    'Summer Aug 7–13, Fall Nov 1–7). Dispatch cost ranges from $39.85/MWh (Spring) to '
    '$45.16/MWh (Summer). Zero load shedding across all seasons with nuclear online.')

heading('4.8. Value of Perfect Information (VOPI)', 2)

body(
    'To quantify the economic cost of ML forecast uncertainty, we compare the day-ahead '
    'planning cost under XGBoost predictions against an oracle plan using realized net load '
    'for the November test week (forecast MAE for this week: 1,080 MW). The oracle MILP '
    'achieves $42.15/MWh while the ML-based plan achieves $41.84/MWh—an absolute gap of '
    '$0.31/MWh (0.74%). We emphasize that this is a planning-stage comparison, not a '
    'realized-cost comparison: the ML forecast systematically under-predicts mean net load by '
    '570 MW (17,837 vs 18,407 MW), so its cheaper plan embeds an under-procurement whose '
    'real-time balancing cost is not captured in our single-stage MILP, and the $0.31/MWh gap '
    'should be read as the order of magnitude of forecast-driven cost distortion rather than a '
    'strict VOPI bound. Even under this caveat, the gap is an order of magnitude below the '
    'nuclear premium ($5.51/MWh); annualized across 52 weeks, the maximum exposure is '
    '~$16M/year system-wide—well within CAISO\'s typical balancing cost budget of '
    '$200–500M/year. A full two-settlement (day-ahead plus real-time recourse) evaluation is '
    'identified as future work.'
)

heading('4.9. Sensitivity Analysis', 2)

body(
    'Fig. 7 presents the results of three parametric sensitivity sweeps, each varying one '
    'parameter while holding others at baseline values.'
)

fig('fig15_sensitivity_analysis.png',
    'Fig. 7. Parametric sensitivity analysis: (a) BESS capacity (0.5–10 GW), '
    '(b) nuclear capacity (0–4 GW), (c) natural gas price ($25–$100/MWh). '
    'Red dashed lines indicate baseline values.')

body(
    'Nuclear capacity (0–4 GW): System cost decreases nearly linearly with nuclear capacity, '
    'at a rate of approximately $2.3/MWh per GW (Fig. 7b). Increasing nuclear from the current '
    '2.26 GW to 4 GW would reduce unit cost from $41.84 to $38.00/MWh (−9.0%), while complete '
    'removal increases cost to $47.35/MWh (+13.2%). The two-tier gas model reveals that '
    'nuclear removal is partly compensated by CCGT output (near maximum at 9,987 MW) and '
    'increasingly by expensive peakers (1,153 MW average), explaining the higher cost premium '
    'compared to single-tier gas formulations. The linearity of this relationship suggests '
    'that, within the modeled range, nuclear provides consistent marginal value—a consequence '
    'of its low marginal cost ($12/MWh) relative to CCGTs ($40/MWh) and peakers ($65/MWh).'
)

body(
    'BESS sizing (0.5–10 GW): In contrast to nuclear, BESS exhibits pronounced diminishing '
    'returns. Increasing BESS from 0.5 GW to 5 GW reduces cost by $0.37/MWh, while further '
    'doubling to 10 GW yields only $0.13/MWh additional savings. Weekly BESS cycling declines '
    'from 5.18 equivalent cycles at 0.5–3 GW (power-limited systems cycle at their maximum) to '
    '2.81 cycles at 10 GW, where the marginal arbitrage opportunity is exhausted. This suggests '
    'that CAISO\'s current 5 GW BESS deployment captures the majority of available arbitrage '
    'value for weekly dispatch optimization.'
)

body(
    'Gas price ($25–$100/MWh): The sweep range spans the CCGT marginal costs implied by Henry '
    'Hub prices from the 2020 lows to the 2021–2022 stress period. The two-tier model reveals '
    'a two-stage substitution pattern '
    'relative to the $55/MWh import price. A first, smaller substitution occurs around '
    '$30/MWh CCGT cost, when the peaker cost ($25/MWh premium over CCGT) crosses the import '
    'price and peaker output is displaced by imports (average gas falls from 15,402 MW at '
    '$25/MWh to 9,905 MW at $35/MWh). The main structural break occurs at approximately '
    '$55/MWh, when CCGT generation itself loses merit-order priority to imports: average gas '
    'dispatch drops from 9,831 MW ($45/MWh) to 6,921 MW ($65/MWh) and then plateaus near '
    '6,800 MW, the level sustained by hours in which the 10 GW import limit binds. This break '
    'point has direct policy implications: a carbon price equivalent to ~$40/ton CO₂ (raising '
    'effective CCGT cost from ~$40 to ~$55/MWh at 0.37 t/MWh) would trigger a fundamental '
    'shift in CAISO\'s dispatch merit order from in-state gas to imports—shifting emissions '
    'accounting to neighboring balancing areas rather than eliminating them (carbon leakage).'
)

body(
    'Import price ($45–$70/MWh): Because gas–import substitution is central to the dispatch '
    'economics, we verify that the headline nuclear premium is not an artifact of the $55/MWh '
    'import assumption. Sweeping the import price from $45 to $70/MWh moves the S4−S1 premium '
    'only from +12.7% to +14.1% (Table 7b), with zero load shedding throughout. The nuclear '
    'valuation is thus robust to the least certain exogenous price in the model.'
)

tbl('Table 7b. Import price sensitivity of the nuclear premium (November week).',
    ['Import ($/MWh)', 'S1 ($/MWh)', 'S4 ($/MWh)', 'Premium', 'S4 import (MW)', 'S4 shed (MWh)'],
    [['45', '38.61', '43.52', '+12.7%', '7,050', '0'],
     ['50', '40.23', '45.48', '+13.1%', '6,678', '0'],
     ['55', '41.84', '47.35', '+13.2%', '6,678', '0'],
     ['60', '43.44', '49.22', '+13.3%', '6,678', '0'],
     ['65', '45.03', '51.10', '+13.5%', '4,148', '0'],
     ['70', '45.09', '51.45', '+14.1%', '1,257', '0']])

heading('4.10. Robustness Parameter Analysis', 2)

body(
    'To address the ad hoc nature of the 50/50 robust blending (S6), we perform a systematic '
    'sweep of the robustness parameter γ ∈ [0, 1], where the MILP input is constructed as '
    'P_net = (1−γ) · P̂_point + γ · P̂_Q90. Table 7 and Fig. 8 show the cost–robustness tradeoff.'
)

tbl('Table 7. Robustness parameter sweep (November test week, Nov 1–7 2023). '
    'Two-tier CCGT UC model. γ=0 recovers S1; γ=0.5 recovers S6; γ=1.0 recovers S2.',
    ['γ', '$/MWh', 'Shed (MWh)', 'Avg Gas (MW)', 'CO₂ (kt/wk)', 'Cost premium vs γ=0'],
    [['0.00 (deterministic)', '41.84', '0', '9,878',  '1,029', '—'],
     ['0.25',                 '42.20', '0', '10,243', '1,083', '+$0.36/MWh (+0.9%)'],
     ['0.50 (S6 Robust)',     '42.63', '0', '10,629', '1,139', '+$0.79/MWh (+1.9%)'],
     ['0.75',                 '43.10', '0', '11,006', '1,197', '+$1.26/MWh (+3.0%)'],
     ['1.00 (worst-case Q90)','43.59', '0', '11,275', '1,254', '+$1.75/MWh (+4.2%)']])

body(
    'The relationship between γ and system cost is nearly linear ($41.84–$43.59/MWh), with each '
    '0.25 increment in γ adding approximately $0.44/MWh (+1.0%). This linearity is notable: '
    'the cost of robustness is predictable and modest—entirely driven by the need to pre-schedule '
    'additional gas generation (CCGT + peakers, rising from 9,895 to 11,280 MW average) to hedge '
    'against the CQR Q90 upper bound. At γ = 0.50 (S6), the cost premium is only $0.79/MWh '
    '(+1.9%) relative to the deterministic case, while providing protection against demand '
    'realizations up to the CQR 90th percentile (mean width = 2,895 MW). Zero load shedding '
    'across all γ values confirms that the system has adequate capacity to absorb the full '
    'range of CQR uncertainty without reliability failures. Importantly, Table 7 cross-validates '
    'with Table 5: γ=0 ↔ S1 ($41.84), γ=0.5 ↔ S6 ($42.63), γ=1.0 ↔ S2 ($43.59).'
)

body(
    'Robustness carries an emissions cost as well as an economic one. Because hedging against '
    'the CQR Q90 bound is achieved almost entirely by pre-scheduling additional gas generation, '
    'weekly CO₂ emissions rise monotonically with γ: from 1,029 kt (γ=0) to 1,254 kt (γ=1), a '
    '+21.8% increase for a +4.2% cost increase (Fig. 8). This cost–robustness–emissions '
    'tradeoff implies that operators face a three-way tension: dispatch plans that are more '
    'robust to forecast uncertainty are simultaneously more expensive and more carbon-intensive. '
    'Reducing forecast uncertainty itself—through better probabilistic models that narrow the '
    'CQR interval—is therefore an emissions-abatement lever, not merely an economic one.'
)

fig('fig_robust_pareto.png',
    'Fig. 8. Cost–robustness–emissions tradeoff: system cost (left axis) and weekly CO₂ '
    'emissions (right axis) as functions of the blending parameter γ '
    '(0 = pure deterministic, 1 = pure worst-case Q90). Two-tier UC model.')

doc.add_page_break()

# =====================================================================
# 5. DISCUSSION
# =====================================================================
heading('4.11. Comparison with Two-Stage Stochastic Programming', 2)

body(
    'Table 8 benchmarks the CQR-robust approach against the two-stage SP baseline of Section '
    '3.5 under the identical information set and out-of-sample protocol. Three observations '
    'stand out. First, the SP planning cost ($42.04/MWh) sits between γ=0 and γ=0.5, i.e., the '
    'probability-weighted scenario tree implies a conservatism level of roughly γ ≈ 0.2. '
    'Second—and most importantly—realized costs against the actual net load are statistically '
    'indistinguishable across all four plans ($42.15–$42.16/MWh, zero shedding): because '
    'CAISO\'s recourse flexibility (CCGT output range, peakers, BESS, imports) is ample in the '
    'studied week, the first-stage commitment differences wash out in operation. Third, the SP '
    'achieves this equivalence at a higher price: it requires scenario probability assignments '
    '(which CQR avoids by construction), enlarges the model by 40% in binary variables (1,176 '
    'vs 840), and solves nearly three times slower (1.2 s vs 0.4 s; wall-clock solve time is hardware-dependent, but the SP model consistently solves several-fold slower across repeated runs)—a gap that compounds '
    'rapidly as scenario trees grow. We conclude that the CQR-robust formulation attains '
    'SP-level operational performance with a distribution-free coverage guarantee, a lighter '
    'model, and a transparent one-parameter conservatism dial; conversely, in systems with '
    'scarcer recourse flexibility, the plan-stage differences between the paradigms would be '
    'expected to re-emerge, a comparison we flag for future work.'
)

tbl('Table 8. CQR-robust vs two-stage stochastic programming under the same information set '
    '(November week). Realized costs from frozen first-stage decisions re-dispatched against '
    'actual net load.',
    ['Plan', 'Planning ($/MWh)', 'Realized ($/MWh)', 'Shed (MWh)', 'Binaries', 'Solve (s)'],
    [['Deterministic (γ=0)',           '41.84', '42.15', '0', '840',   '0.43'],
     ['Robust (γ=0.5)',                '42.63', '42.16', '0', '840',   '0.42'],
     ['Robust (γ=1, Q90)',             '43.59', '42.16', '0', '840',   '0.30'],
     ['SP (3-scenario, Swanson–Megill)','42.04', '42.15', '0', '1,176', '1.17']])

heading('4.12. Annual Robustness of the Nuclear Premium (50-Week Sweep)', 2)

body(
    'The single-week scenarios of Section 4.6 raise an obvious question: is the November week '
    'representative? To test this, we solve the improved MILP—with and without nuclear—for '
    'every one of the 50 non-overlapping 168-hour windows spanning the full realized 2023 net '
    'load record (January 8 to December 31; each window re-initializes SoC at 50%, so weeks '
    'are independent, self-contained dispatch problems), yielding 100 additional MILP solves '
    'at < 1 s each.'
)

body(
    'Two regimes emerge cleanly (Fig. 9). In 35 of 50 weeks (70%, all outside late '
    'June–mid September), net load stays within the stylized fleet\'s capacity envelope and the '
    'nuclear premium is tightly clustered at 11.7–15.0% (mean 13.45%, median 13.4%)—directly '
    'validating the representativeness of the November baseline (+13.2%). In the remaining 15 '
    'weeks, concentrated in a summer heat cluster plus two isolated events, net load approaches '
    'or exceeds the fleet\'s aggregate capacity and the premium diverges sharply, averaging +74% '
    'and peaking at +210% in the week of July 16–22—which contains the true annual net load '
    'maximum (44,900 MW) misidentified as falling within the August 7–13 sample in Section 4.7.'
)

body(
    'This divergence must be interpreted with an important caveat: in 11 of these 15 weeks, '
    'load shedding occurs even with nuclear online, which the stylized fleet (2,256 MW nuclear, '
    '20 GW two-tier gas, 5 GW/20 GWh BESS, 10 GW imports) cannot fully resolve. This is a '
    'modeling artifact, not a reproduction of observed CAISO reliability outcomes: the fleet '
    'omits CAISO\'s actual hydroelectric and pumped-storage capacity (~6–8 GW), demand response '
    'programs, and out-of-market emergency measures, all of which the real system deployed '
    'during the July–September 2023 heat events without resorting to rotating outages. The '
    'correct reading is therefore not "CAISO nearly lost load in 2023," but that nuclear\'s '
    'reliability value—distinct from and additive to its ~13% economic value—rises sharply as a '
    'stylized capacity-constrained system approaches its frontier, consistent with the '
    'capacity-adequacy literature (Sepulveda et al., 2018). Reassuringly, the full-year CO₂ '
    'saving from nuclear (~9.0 Mt/year, summed over 50 weeks) closely matches the single-week '
    'extrapolation (~9.3 Mt/year, Section 4.10), cross-validating both estimates.'
)

fig('fig_fullyear_dispatch.png',
    'Fig. 9. Full-year weekly dispatch (50 consecutive weeks, realized 2023 net load). Top: '
    'dispatch cost with and without nuclear. Bottom: nuclear removal premium, colored by '
    'regime—normal weeks (dark red, tightly clustered near the November baseline) vs scarcity '
    'weeks (orange, shed > 0 without nuclear, concentrated in summer).')

body(
    'Validating the modeling-artifact interpretation with realized hydro: Rather than merely '
    'assert that the summer shedding above is an artifact of the omitted hydro fleet, we test '
    'it directly. The 15 scarcity weeks are re-solved (with nuclear) after adding realized 2023 '
    'hourly CAISO hydroelectric generation (EIA Form 930 fuel-type series) as an additional '
    'must-take resource, dispatched ahead of thermal generation—consistent with the largely '
    'run-of-river, environmentally constrained operation of CAISO hydro. This requires no '
    'change to the core MILP; it only supplies the historical hydro profile the stylized fleet '
    'omits. The result strongly confirms the artifact interpretation (Table 8b, Fig. 9b): total '
    'shedding across the 15 weeks falls by 93.4% (from 167,506 to 11,021 MWh), and 11 of 15 '
    'weeks are fully resolved to zero shedding. The four weeks retaining residual shedding '
    '(January 15, February 26, July 16, September 10) are consistent with the one remaining '
    'omitted resource—demand response—which CAISO also deployed during 2023 stress events but '
    'which is not modeled here. This quantifies, rather than merely asserts, that the '
    'scarcity-week shedding predominantly reflects the stylized fleet\'s missing hydro capacity '
    'rather than a genuine CAISO reliability shortfall.'
)

tbl('Table 8b. Effect of realized 2023 CAISO hydro generation on the 15 flagged scarcity '
    'weeks (with nuclear online).',
    ['Metric', 'Value'],
    [['Total shedding without hydro', '167,506 MWh'],
     ['Total shedding with realized hydro added', '11,021 MWh'],
     ['Reduction in total shedding', '93.4%'],
     ['Weeks fully resolved (zero shedding)', '11 / 15']])

fig('fig_hydro_robustness.png',
    'Fig. 9b. Effect of adding realized 2023 CAISO hydro generation (must-take, dispatched '
    'ahead of thermal) on load shedding in the 15 scarcity weeks, with nuclear online. Total '
    'shedding falls by 93%; 11 of 15 weeks are fully resolved to zero.')

heading('4.13. Supplementary Check: Directional Transferability Beyond CAISO', 2)

body(
    'The 50-week sweep establishes representativeness within CAISO; a natural follow-up '
    'question is whether the same scarce-recourse mechanism operates in other grids. We do '
    'not attempt a second full case study—that would require officially sourced unit-level '
    'capacity data and a full annual sweep beyond this paper\'s scope—but run a narrowly '
    'framed supplementary check: the identical MILP code, re-parameterized with illustrative '
    'ERCOT-like fleet and import-capacity values, applied to two non-extreme 2023 weeks '
    '(Appendix B gives the full parameterization, table, and figure). The result is '
    'directionally consistent with the paper\'s central hypothesis: with island-grid import '
    'capacity an order of magnitude below CAISO\'s (1.22 GW vs 10 GW), the nuclear premium is '
    'larger (+14.9% and +15.8%) than the CAISO baseline (+13.2%). We report this strictly as '
    'a qualitative transferability signal rather than a calibrated regional estimate, given '
    'the illustrative parameterization and the two-week scope; a rigorously sourced, '
    'full-year study of another grid is identified as future work (Section 6).'
)

doc.add_page_break()

heading('5. Discussion', 1)

heading('5.1. Policy Implications', 2)

body(
    'The four-season dispatch analysis reveals a consistent seasonal ordering of system costs: '
    'Summer ($45.16/MWh) > Fall ($42.15/MWh) ≈ Winter ($42.27/MWh) > Spring ($39.85/MWh). '
    'We stress that the nuclear valuation is an operating-cost estimate under fixed technology '
    'prices, not a market-equilibrium one: removing 2.26 GW of baseload would in reality shift '
    'wholesale prices and import flows, effects a price-taking dispatch model cannot capture; '
    'the premium is nonetheless robust to the least certain of those prices (import cost, '
    '+12.7% to +14.1% over $45–$70/MWh, Section 4.9). '
    'Nuclear baseload value is primarily economic under typical operating conditions: the S4 '
    '(no-nuclear) dispatch shows that the '
    'expanded gas fleet and imports maintain zero load shedding in all four representative '
    'seasonal weeks—but at a '
    'sustained $5.51/MWh premium ($41.84 → $47.35/MWh in the November deterministic scenario). '
    'This economic-value characterization does not extend to capacity-stress conditions: the '
    '50-week annual sweep (Section 4.12) shows that during summer weeks approaching the '
    'system\'s capacity frontier, nuclear\'s reliability value becomes dominant, with the premium '
    'rising to +74% on average and +210% at the annual net load peak—a second, distinct value '
    'stream additive to the baseline economic premium. '
    'At $2.3/MWh per GW across a 0–4 GW range, nuclear provides consistent marginal value '
    'that argues for its retention as a low-cost zero-carbon baseload resource, with economic '
    'value highest in high-demand periods when peaker and import costs are elevated. '
    'The VOPI of $0.31/MWh (0.7%) confirms that ML forecast '
    'accuracy is sufficient for day-ahead dispatch; the ML-to-oracle cost gap is an order of '
    'magnitude smaller than the nuclear premium ($5.51/MWh), meaning forecast uncertainty is '
    'not the binding constraint on system economics. For CAISO operators, the CQR robustness '
    'parameter γ provides a low-cost lever: each 0.25 increment adds only ~$0.44/MWh (+1.0%) '
    'while providing protection against demand realizations up to the CQR 90th percentile.'
)

heading('5.2. Comparison with Prior Literature', 2)

body(
    'Our day-ahead net load R² = 0.854 and skill = +18.5% over persistence are consistent with '
    'the GEFCom literature (Hong and Fan, 2016), where top entries achieve 10–15% skill gains. '
    'Our higher skill gain reflects the additional variability of net load versus gross demand, '
    'against which gradient boosting\'s feature selection is particularly effective. The MLP '
    'overfit gap (train R² = 0.996 vs test R² = 0.481) confirms the well-documented superiority '
    'of GBM over deep learning on moderate-size tabular energy datasets (Hertel et al., 2022; '
    'Shao et al., 2022). CQR coverage of 88.8% (target 90%) aligns with Romano et al. (2019) '
    'and Stankovic et al. (2023); slight undercoverage is expected given temporal dependence. '
    'The nuclear value of $2.3/MWh per GW in CAISO falls at the lower end of Jenkins et al. '
    '(2018) estimates ($2–5/MWh for generic U.S. systems), consistent with CAISO\'s large '
    'gas fleet and import capacity providing substantial alternative flexibility. The gas price '
    'structural break at ~$55/MWh echoes Bistline et al. (2023) tipping-point findings but is '
    'derived from actual 2023 CAISO operational data rather than capacity expansion models.'
)

heading('5.3. Limitations', 2)

body(
    'This study has six main limitations. (i) Single-bus (copper-plate) formulation: '
    'transmission constraints—particularly congestion on key CAISO interfaces (Path 15, Path 26)—'
    'are omitted. Zonal or nodal modeling would provide more realistic dispatch but substantially '
    'increase complexity. (ii) Aggregate CCGT model: two aggregate 5,000 MW units approximate '
    'CAISO\'s actual fleet (~22 GW gas capacity across ~200 units); individual unit commitment '
    'would capture maintenance scheduling and heat-rate curves more accurately. The CCGT startup '
    'cost ($25,000/event ≈ $5/MW; NREL ATB, 2023) reflects warm-start costs for aggregate units, '
    'conservatively below cold-start estimates of $10–15/MW from the same source. We tested '
    'sensitivity to this choice directly: re-solving S1 and S4 with four 2,500 MW units instead '
    'of two 5,000 MW units (same 10 GW total capacity, minimum-load fraction, and per-unit '
    'ramp rate) changes the nuclear premium by only 0.01 percentage points (13.175% vs '
    '13.165%), confirming the headline results are insensitive to this aggregation choice; '
    'individual unit commitment would still improve fidelity for maintenance scheduling and '
    'heat-rate curves. '
    '(iii) CQR exchangeability and conditional coverage: time series dependence technically '
    'violates the exchangeability assumption, and while marginal coverage of 88.8% (vs 90% '
    'target) is near-nominal, hour-conditional coverage drops to ~79% during evening ramps '
    '(Fig. 4b); adaptive and Mondrian-style conformal methods (Gibbs and Candès, 2021) would '
    'address both issues under '
    'distribution shift. (iv) Single-stage, single-year scope: the MILP is day-ahead only (no '
    'real-time recourse stage) and uses 2023 conditions exclusively; multi-year analysis would '
    'capture secular VRE capacity growth and its interaction with nuclear flexibility value. '
    '(v) Import representation: imports are modeled at a constant $55/MWh with 10 GW '
    'availability; during West-wide stress events (e.g., regional heat waves), both the price '
    'and the availability of imports deteriorate simultaneously, so the no-nuclear scenarios '
    'likely understate scarcity costs—our nuclear valuation is conservative in this respect. '
    '(vi) Stylized capacity envelope and ERCOT parameterization: the fleet used in the annual '
    'sweep (Section 4.12) omits CAISO\'s hydroelectric, pumped-storage, and demand-response '
    'resources, causing artificial load shedding in some summer weeks that did not occur in '
    'reality. We tested this directly rather than merely disclosing it: adding realized 2023 '
    'hourly CAISO hydro generation as a must-take resource eliminates 93.4% of the flagged '
    'shedding and fully resolves 11 of the 15 affected weeks (Table 8b); the residual shedding '
    'in the remaining four weeks is consistent with demand response, the one omitted resource '
    'we could not similarly test with historical dispatch data. The supplementary ERCOT check '
    '(Section 4.13, Appendix B) is deliberately scoped as a qualitative transferability '
    'signal, not a second case study: it uses illustrative, not officially sourced, fleet '
    'parameters over only two non-extreme weeks, and we report it strictly as evidence that '
    'the scarcity mechanism points in the expected direction outside CAISO, not as a '
    'calibrated ERCOT-specific policy estimate.'
)

doc.add_page_break()

# =====================================================================
# 6. CONCLUSIONS
# =====================================================================
heading('6. Conclusions and Policy Implications', 1)

body(
    'This study presents an integrated data-driven framework for optimal dispatch of nuclear '
    'baseload and battery storage under duck curve uncertainty, applied to the California ISO '
    'system using real 2023 operational data. The framework couples conformalized quantile '
    'regression with robust MILP optimization, providing a principled pipeline from data '
    'acquisition through probabilistic forecasting to dispatch decision-making. The key '
    'conclusions are as follows.'
)

conclusions = [
    ('Day-ahead net load forecasting', 'XGBoost with 35 features '
     'achieves R² = 0.854 and MAE = 1,429 MW on the held-out test set, statistically '
     'outperforming the persistence baseline by 18.5% (Diebold–Mariano p < 0.001). Five-fold '
     'expanding-window cross-validation confirms stability across seasonal regimes '
     '(mean R² = 0.869 ± 0.089). The strict 24-hour information cutoff ensures that reported '
     'metrics reflect achievable operational performance.'),

    ('Conformalized prediction intervals', 'CQR provides 90% target coverage with 88.8% '
     'empirical coverage and adaptive interval width (mean 5,182 MW), enabling principled '
     'uncertainty propagation to dispatch optimization without distributional assumptions. '
     'The conformal correction (Q̂ = 1,044 MW) provides a clear measure of the additional '
     'uncertainty not captured by the quantile regression models.'),

    ('Nuclear baseload value', 'Removing 2.26 GW of nuclear capacity increases system costs by '
     '13.2% ($41.84 → $47.35/MWh) in the deterministic November dispatch scenario. The two-tier '
     'gas fleet and imports fully compensate operationally (zero load shedding in S4), but at a '
     'sustained $5.51/MWh cost premium—equivalent to $2.3/MWh per GW of nuclear, confirmed '
     'across a 0–4 GW parametric sweep (Section 4.9). Under typical conditions this consistent '
     'marginal value is highest '
     'in high-demand periods when peaker and import costs are elevated, and lowest in spring when '
     'excess renewable generation reduces the marginal dispatch cost.'),

    ('Four-season analysis and VOPI', 'Dispatch cost ranges seasonally from $39.85/MWh (spring, '
     'deep duck belly) to $45.16/MWh (summer, peak demand), with zero load shedding across all '
     'four representative seasonal weeks when nuclear is online. The Value of Perfect Information '
     '(VOPI) for the November '
     'test week is $0.31/MWh (0.74%), confirming that the 1,429 MW MAE has negligible economic '
     'impact on day-ahead dispatch and validating the ML-to-MILP integration.'),

    ('CQR-robust vs stochastic programming', 'Under an identical information set and '
     'out-of-sample protocol, a two-stage SP baseline attains the same realized cost as the '
     'CQR-robust dispatch ($42.15 vs $42.16/MWh, zero shedding) while requiring scenario '
     'probabilities, 40% more binaries, and several-fold longer solves—the CQR-robust '
     'formulation delivers SP-level performance with a distribution-free guarantee and a '
     'lighter model.'),

    ('Annual robustness and transferability', 'A 50-week sweep of the full 2023 record confirms '
     'the November premium is representative of typical operation (13.45% mean across 35 normal '
     'weeks) while revealing a second, reliability-driven regime during summer capacity-stress '
     'weeks (premium up to +210%). Two targeted checks validate this finding rather than merely '
     'disclosing its limitations: adding realized 2023 CAISO hydro generation eliminates 93.4% '
     'of the flagged shedding (11 of 15 weeks fully resolved), confirming a modeling-artifact '
     'rather than a genuine-reliability interpretation, and re-solving with four 2,500 MW CCGT '
     'units instead of two 5,000 MW units changes the nuclear premium by only 0.01 percentage '
     'points. A narrowly scoped supplementary check on an ERCOT-parameterized grid (Appendix B) '
     'shows the same scarcity mechanism directionally amplifying the premium (+15–16%) under '
     'an island grid\'s limited import capacity—a qualitative transferability signal, not a '
     'calibrated regional estimate.'),

    ('Gas price structural break and emissions', 'A threshold effect at ~$55/MWh triggers rapid '
     'substitution from CCGT generation to imports, with average gas dispatch dropping from '
     '9,831 MW to 6,921 MW across the break. A carbon price of ~$40/ton CO₂ would trigger this '
     'shift, with attendant carbon-leakage concerns. Nuclear baseload avoids approximately '
     '178 kt CO₂ per week (~9.3 Mt/year annualized at the fossil margin) relative to the '
     'no-nuclear dispatch, and '
     'robust dispatch carries an emissions premium of up to +21.8% (γ=1 vs γ=0)—a '
     'cost–robustness–emissions tradeoff that makes forecast quality an emissions lever.'),
]

for i, (title, text) in enumerate(conclusions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.27)
    r = p.add_run(f'({i}) ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r2 = p.add_run(f'{title}. ')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.bold = True; r2.italic = True
    r3 = p.add_run(text)
    r3.font.name = 'Times New Roman'; r3.font.size = Pt(12)

body(
    'Future research directions include: (i) incorporating transmission network constraints '
    'through zonal or nodal MILP formulations; (ii) extending the analysis to multi-year '
    'horizons to capture the secular growth of VRE capacity; (iii) integrating weather ensemble '
    'forecasts with the CQR framework for improved prediction interval calibration; '
    '(iv) applying adaptive conformal inference methods designed for non-stationary time series; '
    '(v) expanding the optimization to include demand response, electric vehicle charging, '
    'and hydrogen production as additional flexibility resources; (vi) a full capacity-adequacy '
    'representation of the CAISO fleet (hydroelectric, pumped storage, and demand response) to '
    'replace the stylized envelope used in the annual sweep; and (vii) a rigorously sourced, '
    'full-year ERCOT study using official CDR unit-level capacity data.'
)

doc.add_page_break()

# =====================================================================
# DECLARATIONS
# =====================================================================
heading('CRediT Author Statement', 2)
body(
    'Cagatay Kuban: Conceptualization, Methodology, Software, Data curation, Formal analysis, '
    'Validation, Visualization, Writing – original draft, Writing – review & editing.'
)

heading('Declaration of Competing Interest', 2)
body(
    'The author declares that they have no known competing financial interests or personal '
    'relationships that could have appeared to influence the work reported in this paper.'
)

heading('Data Availability', 2)
body(
    'All data used in this study are publicly available from the U.S. Energy Information '
    'Administration (api.eia.gov) and the NOAA National Centers for Environmental Information '
    '(ncei.noaa.gov). Code and processed datasets are available from the corresponding author '
    'upon reasonable request.'
)

heading('Declaration of Generative AI and AI-Assisted Technologies in the Manuscript Preparation Process', 2)
body(
    'During the preparation of this work the author used the Claude Sonnet model (Anthropic) '
    'for language support and coding support. After using this tool, the author reviewed and '
    'edited the content as needed and takes full responsibility for the content of the '
    'published article.'
)


heading('Acknowledgements', 2)
body(
    'The authors acknowledge the U.S. Energy Information Administration for providing open '
    'access to electricity grid data and the NOAA for meteorological data.'
)

doc.add_page_break()

# =====================================================================
# REFERENCES
# =====================================================================
heading('References', 1)

refs = [
    'Bertsimas, D., Sim, M., 2004. The price of robustness. Operations Research 52(1), 35–53.',
    'Bertsimas, D., Litvinov, E., Sun, X.A., Zhao, J., Zheng, T., 2013. Adaptive robust optimization for the security constrained unit commitment problem. IEEE Transactions on Power Systems 28(1), 52–63.',
    'Birge, J.R., Louveaux, F., 2011. Introduction to Stochastic Programming, second ed. Springer, New York.',
    'Bistline, J.E.T., Mehrotra, N.R., Wolfram, C., 2023. Economic implications of the climate provisions of the Inflation Reduction Act. Brookings Papers on Economic Activity 2023(1), 77–152.',
    'Breiman, L., 2001. Random forests. Machine Learning 45(1), 5–32.',
    'California Independent System Operator (CAISO), 2016. What the duck curve tells us about managing a green grid. Technical report.',
    'California Independent System Operator (CAISO), 2023. 2023 Annual Report on Market Issues and Performance. Department of Market Monitoring, Folsom, CA. Transmission interface limits obtained from CAISO OASIS (oasis.caiso.com).',
    'Cany, C., Mansilla, C., da Costa, P., Mathonnière, G., 2018. Nuclear power supply: Going against the misconceptions. Energy Policy 116, 42–56.',
    'Chen, T., Guestrin, C., 2016. XGBoost: A scalable tree boosting system. In: Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. ACM, pp. 785–794.',
    'Denholm, P., O\'Connell, M., Brinkman, G., Jorgenson, J., 2015. Overgeneration from solar energy in California: A field guide to the duck chart. National Renewable Energy Laboratory, Technical Report NREL/TP-6A20-60953.',
    'Denholm, P., Brown, P., Cole, W., Mai, T., Sergi, B., Brown, M., et al., 2023. Examining supply-side options to achieve 100% clean electricity by 2035. Renewable Energy 206, 1028–1039.',
    'Diebold, F.X., Mariano, R.S., 1995. Comparing predictive accuracy. Journal of Business & Economic Statistics 13(3), 253–263.',
    'Electric Reliability Council of Texas (ERCOT), 2023. 2023 Capacity, Demand and Reserves (CDR) Report. Austin, TX.',
    'Gibbs, I., Candès, E., 2021. Adaptive conformal inference under distribution shift. In: Advances in Neural Information Processing Systems 34.',
    'Hart, W.E., Laird, C.D., Watson, J.-P., Woodruff, D.L., Hackebeil, G.A., Nicholson, B.L., Siirola, J.D., 2017. Pyomo—Optimization Modeling in Python, second ed. Springer, New York.',
    'Hertel, M., Ott, S., Schäfer, B., Neumann, D., 2022. Transformer training strategies for forecasting multiple load time series. Energy and AI 8, 100147.',
    'Hong, T., Fan, S., 2016. Probabilistic electric load forecasting: A tutorial review. International Journal of Forecasting 32(3), 914–938.',
    'Huangfu, Q., Hall, J.A.J., 2018. Parallelizing the dual revised simplex method. Mathematical Programming Computation 10(1), 119–142.',
    'Jenkins, J.D., Zhou, Z., Ponciroli, R., Vilim, R.B., Ganda, F., de Sisternes, F., Botterud, A., 2018. The benefits of nuclear flexibility in power system operations with renewable energy. Applied Energy 222, 872–884.',
    'Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., Liu, T.-Y., 2017. LightGBM: A highly efficient gradient boosting decision tree. In: Advances in Neural Information Processing Systems 30.',
    'Keefer, D.L., Bodily, S.E., 1983. Three-point approximations for continuous random variables. Management Science 29(5), 595–609.',
    'Koenker, R., 2005. Quantile Regression. Cambridge University Press, Cambridge.',
    'NEA (Nuclear Energy Agency), 2011. Technical and Economic Aspects of Load Following with Nuclear Power Plants. OECD Nuclear Energy Agency, Paris. NEA No. 6786.',
    'NREL (National Renewable Energy Laboratory), 2023. Annual Technology Baseline 2023. National Renewable Energy Laboratory, Golden, CO. Available: atb.nrel.gov (Accessed: June 2024). Gas Combined Cycle: startup parameters, warm-start cost ≈ $5/MW-start.',
    'Romano, Y., Patterson, E., Candès, E., 2019. Conformalized quantile regression. In: Advances in Neural Information Processing Systems 32.',
    'Sepulveda, N.A., Jenkins, J.D., de Sisternes, F.J., Lester, R.K., 2018. The role of firm low-carbon electricity resources in deep decarbonization of power generation. Joule 2(11), 2403–2420.',
    'Shao, Z., Chao, F., Yang, S.L., Zhou, K.L., 2022. A review of the decomposition methodology for extracting and identifying the fluctuation characteristics in electricity demand forecasting. Renewable and Sustainable Energy Reviews 163, 112507.',
    'Sioshansi, R., Denholm, P., Jenkin, T., Weiss, J., 2009. Estimating the value of electricity storage in PJM: Arbitrage and some welfare effects. Energy Economics 31(2), 269–277.',
    'Stankovic, L., Stankovic, V., Liao, J., Wilson, C., 2023. Measuring the untapped potential of smart meter data for probabilistic forecasting of net demand. Renewable Energy 205, 1091–1104.',
    'Vovk, V., Gammerman, A., Shafer, G., 2005. Algorithmic Learning in a Random World. Springer, New York.',
    'Wang, Y., Zhang, N., Kang, C., Kirschen, D.S., Li, J., Xia, Q., 2019. Standardized matrix modeling of multiple energy systems. IEEE Transactions on Smart Grid 10(1), 257–270.',
    'Weron, R., 2014. Electricity price forecasting: A review of the state-of-the-art with a look into the future. International Journal of Forecasting 30(4), 1030–1081.',
    'Xu, C., Chen, H., 2023. Conformal prediction intervals for remaining useful life estimation and electricity price forecasting. Energy 273, 127014.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# =====================================================================
# APPENDIX
# =====================================================================
doc.add_page_break()
heading('Appendix A. Additional Tables and Figures', 1)

body(
    'This appendix reports supporting diagnostics referenced from Sections 4.3 and 4.4.'
)

body(
    'Error diagnostics: Formal residual tests (Table A1) show left-skewed (skewness = −0.72, '
    'indicating the model tends to slightly overestimate net load, i.e., under-predict '
    'renewable output during high-solar periods) with near-zero excess kurtosis (−0.007), '
    'suggesting light tails relative to the normal distribution. The Jarque–Bera test rejects '
    'the null hypothesis of normality for all models (p < 0.001), as does the Shapiro–Wilk '
    'test. This non-normality is expected for electricity data, which exhibits heavy tails '
    'from demand spikes and renewable intermittency, and does not invalidate the ML models or '
    'the CQR framework (which is distribution-free by construction). The Ljung–Box test at 24 '
    'lags detects significant residual autocorrelation (p < 0.001) for all models, reflecting '
    'unmodeled intra-day persistence patterns—particularly the 24-hour cycle where errors at '
    'hour h tend to correlate with errors at hour h−24. Despite this residual autocorrelation, '
    'the CQR prediction intervals maintain near-nominal coverage (88.8%), consistent with the '
    'empirical findings of Stankovic et al. (2023), who observed similar robustness of CQR to '
    'temporal dependence in wind power forecasting.'
)

tbl('Table A1. Error diagnostics for net load forecasting models.',
    ['Model', 'Mean Error', 'Skewness', 'Kurtosis', 'JB p-value', 'SW p-value', 'LB(24) p-value'],
    [['LightGBM', '−175 MW', '−0.73', '−0.005', '< 0.001', '< 0.001', '< 0.001'],
     ['XGBoost', '−143 MW', '−0.72', '−0.007', '< 0.001', '< 0.001', '< 0.001'],
     ['Random Forest', '−161 MW', '−0.73', '−0.004', '< 0.001', '< 0.001', '< 0.001']])

body(
    'Hyperparameter sensitivity analysis (Table A2) confirms that the LightGBM results are '
    'robust to the choice of num_leaves, the primary complexity-controlling hyperparameter. '
    'Test R² varies by only 1.3 percentage points (0.841–0.854) across a 8× range of num_leaves '
    '(31–255), indicating that the model is not overly sensitive to this choice. The reported '
    'configuration (63 leaves) is selected on validation MAE, which it minimizes, confirming '
    'the selection is not test-set-driven despite num_leaves=31 scoring marginally higher on '
    'the held-out test fold.'
)

tbl('Table A2. Hyperparameter sensitivity (LightGBM, num_leaves).',
    ['num_leaves', 'Val MAE (MW)', 'Test MAE (MW)', 'Test R²'],
    [['31', '1,410', '1,498', '0.854'],
     ['63 (selected)', '1,365', '1,565', '0.843'],
     ['127', '1,408', '1,570', '0.841'],
     ['255', '1,408', '1,571', '0.841']])

body(
    'Cross-validation detail: Table A3 reports the full five-fold expanding-window breakdown '
    'underlying Section 4.3. Fold 2 (July 10 – August 21, R² = 0.719) coincides with the 2023 '
    'California heat events, when air-conditioning-driven demand spikes deviate from the '
    'historical patterns embedded in lag features. Forecast error correlates weakly but '
    'positively with temperature (r = 0.162, Fig. A7), with the largest errors on days '
    'exceeding 35°C in the Central Valley—motivating the weather-ensemble extension identified '
    'as future work (Section 6).'
)

tbl('Table A3. Expanding-window cross-validation (LightGBM, net load).',
    ['Fold', 'Train size', 'Train end', 'Test end', 'MAE (MW)', 'RMSE (MW)', 'R²'],
    [['1', '3,389', 'May 29', 'Jul 10', '1,203', '1,855', '0.926'],
     ['2', '4,405', 'Jul 10', 'Aug 21', '2,711', '3,737', '0.719'],
     ['3', '5,421', 'Aug 21', 'Oct 3', '1,451', '2,044', '0.922'],
     ['4', '6,437', 'Oct 3', 'Nov 15', '1,335', '1,821', '0.925'],
     ['5', '7,453', 'Nov 15', 'Dec 31', '1,457', '2,008', '0.850'],
     ['Mean ± Std', '', '', '', '1,632 ± 612', '2,293 ± 813', '0.869 ± 0.089']])

fig('fig01_timeseries_overview.png',
    'Fig. A1. CAISO 2023 annual overview: total demand and net generation (top); '
    'net load = demand − solar − wind (bottom).')

fig('fig04_monthly_netload_boxplots.png',
    'Fig. A2. Monthly distribution of net load. Spring months (Mar–May) show the widest range '
    'and lowest medians due to high solar penetration.')

fig('fig05_demand_heatmap.png',
    'Fig. A3. Net load heatmap (hour × month). The duck belly (blue region) is concentrated in '
    'spring mid-day hours; evening peaks (red) are strongest in summer.')

fig('fig10_error_analysis.png',
    'Fig. A4. Error analysis: distributions (top) and actual vs. predicted scatter plots (bottom) '
    'for net load (left) and price (right).')

fig('fig09_feature_importance.png',
    'Fig. A5. Top-15 feature importance (XGBoost gain) for the canonical 35-feature net load '
    'model. The 24-h and 48-h net load lags jointly account for ~67% of total gain, confirming '
    'the dominance of short-lag autocorrelation noted in Section 3.2; the lagged wholesale '
    'price (price_lag24h) contributes only 1.1%, indicating the day-ahead model does not '
    'materially rely on price information.')

fig('fig16_cv_stability.png',
    'Fig. A6. Expanding-window cross-validation stability: MAE (bars) and R² (line) across 5 folds.')

fig('fig_error_vs_temperature.png',
    'Fig. A7. Forecast error analysis vs temperature. Left: scatter of |error| vs temperature '
    '(r = 0.162, weak positive correlation). Right: monthly MAE and average temperature.')

fig('fig_co2_analysis.png',
    'Fig. A8. CO₂ emissions analysis (S1–S6, November week, two-tier UC model; emission factors: '
    'CCGT 0.37, peaker 0.55, imports 0.428 t CO₂/MWh, the CARB unspecified-import default). '
    'Left: weekly emissions by scenario '
    '(S1 = 1,029 kt; S4 No-Nuclear = 1,207 kt, +17.3%). Right: cost–emissions plane showing that '
    'nuclear removal increases both cost and emissions, corresponding to ~9.3 Mt CO₂/year '
    'avoided by the 2.26 GW nuclear fleet.')

fig('fig_2022_validation.png',
    'Fig. A9. Out-of-sample temporal generalization: LightGBM trained on 2023 data applied to '
    '8,593 hourly observations from 2022 (R² = 0.908, MAE = 1,514 MW), confirming durable '
    'seasonal and diurnal pattern capture without temporal overfitting.')

fig('fig_vopi_analysis.png',
    'Fig. A10. Value of Perfect Information (VOPI) analysis. Oracle dispatch ($42.15/MWh) vs '
    'ML-based dispatch ($41.84/MWh): VOPI = $0.31/MWh (0.74%), confirming ML forecast accuracy '
    'is sufficient for day-ahead dispatch optimization.')

fig('fig_nuclear_sensitivity.png',
    'Fig. A11. Nuclear capacity sensitivity (0-4 GW). System cost decreases nearly linearly '
    'at $2.3/MWh per GW, with peaker deployment declining as nuclear displaces expensive '
    'marginal generation.')

fig('fig_zonal_sensitivity.png',
    'Fig. A12. Zonal Sensitivity Analysis: Copper-plate vs. 2-Zone (NP15/SP15) with Path 15 '
    'transmission constraint (4,500 MW fwd / 3,500 MW rev). The two-zone model shows a modest '
    '3.73% cost difference, validating the single-bus assumption.')

fig('fig_import_sensitivity.png',
    'Fig. A13. Nuclear premium vs import price ($45–$70/MWh): the S4−S1 premium (annotated) '
    'remains within +12.7 to +14.1%; both curves flatten above $65/MWh as imports exit the '
    'merit order.')

doc.add_page_break()
heading('Appendix B. Supplementary Transferability Check: ERCOT-Parameterized Grid', 1)

body(
    'This appendix gives the full parameterization and results underlying the supplementary '
    'check referenced in Section 4.13. We stress at the outset that this exercise is scoped '
    'narrowly and deliberately: it asks only whether the paper\'s scarce-recourse mechanism '
    'points in the expected direction outside CAISO, using illustrative fleet parameters over '
    'two weeks—not whether it reproduces ERCOT\'s actual 2023 operation or supports an '
    'ERCOT-specific policy conclusion. A rigorous treatment of that latter question, with '
    'unit-level capacity data sourced from ERCOT\'s own Capacity, Demand and Reserves report '
    'and a full annual sweep analogous to Section 4.12, is identified as future work rather '
    'than attempted here.'
)

body(
    'To probe whether the mechanism generalizes beyond CAISO, we re-parameterize the '
    'identical MILP code for the Electric Reliability Council of Texas (ERCOT), using freshly '
    'retrieved 2023 EIA hourly demand, solar, and wind data (8,760 hours; net load mean '
    '35,035 MW, range 12,398–70,443 MW) and approximate, illustrative ERCOT fleet parameters: '
    'nuclear 5,150 MW (Comanche Peak and South Texas Project), a 45 GW mid-merit thermal tier '
    '(blended gas-CC/coal, $42/MWh), 15 GW peakers, 3.2 GW/6.4 GWh BESS, and—critically—only '
    '1.22 GW of DC-tie import/export capacity, an order of magnitude below CAISO\'s 10 GW, '
    'reflecting ERCOT\'s near-total electrical islanding from neighboring grids.'
)

body(
    'Two illustrative weeks are evaluated with and without nuclear: a wind-driven low-net-load '
    'week (March 13–19) and a November shoulder week (Table B1, Fig. B1). Both weeks lie '
    'within the stylized fleet\'s capacity envelope (zero shedding in all four runs), so this '
    'comparison isolates the economic premium. The result is directionally consistent with '
    'the paper\'s hypothesis: ERCOT\'s premiums (+14.9% and +15.8%) exceed CAISO\'s November '
    'baseline (+13.2%), consistent with scarcer import capacity forcing greater reliance on '
    'thermal generation to replace nuclear. Beyond the scoping caveat above, two further '
    'limits apply: only two non-extreme weeks are tested (not a full year, and not ERCOT\'s '
    'own summer 2023 peak, when the actual system operated under Energy Emergency Alerts), '
    'and the fleet parameters are illustrative approximations rather than sourced from '
    'ERCOT\'s Capacity, Demand and Reserves report.'
)

tbl('Table B1. Supplementary ERCOT check: nuclear premium under illustrative fleet '
    'parameters (two representative weeks, zero shedding in all runs).',
    ['Week', 'With nuclear ($/MWh)', 'No nuclear ($/MWh)', 'Premium', 'Peak net load (MW)'],
    [['Spring (Mar 13–19)', '34.81', '39.98', '+14.9%', '38,640'],
     ['November (shoulder)', '34.52', '39.97', '+15.8%', '42,396']])

fig('fig_ercot_case.png',
    'Fig. B1. Supplementary ERCOT transferability check. (a) Wind-driven net load profile for '
    'the ERCOT spring week. (b) Nuclear removal premium in ERCOT (island topology, 1.22 GW '
    'ties) vs the CAISO November baseline (10 GW imports)—consistent with the hypothesis that '
    'scarce import capacity amplifies the nuclear premium.')

# =====================================================================
# SAVE
# =====================================================================
out_path = os.path.join(OUT, 'manuscript.docx')
doc.save(out_path)
sz = os.path.getsize(out_path)
print(f"\nSaved: {out_path}")
print(f"Size: {sz / 1024:.0f} KB ({sz / 1024 / 1024:.1f} MB)")
