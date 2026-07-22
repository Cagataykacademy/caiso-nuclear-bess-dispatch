"""Generate the Highlights file as a separate submission item (Applied Energy requires
this uploaded as its own file, with 'highlights' in the filename; 3-5 bullets, <=85 chars each)."""
import os
from docx import Document
from docx.shared import Pt

HIGHLIGHTS = [
    'True CAISO net load from EIA fuel-type data reveals a 55 GW duck curve depth.',
    'XGBoost day-ahead forecast: R2 = 0.854, +18.5% skill over persistence.',
    'CQR-robust MILP matches stochastic programming with 40% fewer binary variables.',
    'Nuclear premium is 13% in typical weeks but reaches 210% under capacity stress.',
    'Robust dispatch trades +4.2% cost for +21.8% CO2, an emissions-abatement lever.',
]

for h in HIGHLIGHTS:
    assert len(h) <= 85, f"Highlight exceeds 85 chars ({len(h)}): {h}"

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

for h in HIGHLIGHTS:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(h)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'paper')
path = os.path.join(OUT, 'highlights.docx')
doc.save(path)
print(f'Saved: {path}')
